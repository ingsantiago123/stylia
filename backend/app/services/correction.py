"""
Servicio de Corrección (Etapa D).
Pipeline: LanguageTool (ortografía/gramática) → ChatGPT (estilo/claridad).

Para Ruta 1 (DOCX-first): corrige párrafos directamente del DOCX,
no de bloques del PDF, para evitar fragmentación y mayúsculas erróneas.
"""

import json
import logging
import tempfile
from collections.abc import Callable
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from pathlib import Path

import httpx
from docx import Document as DocxDocument
from tenacity import retry, stop_after_attempt, wait_fixed, retry_if_exception_type

from app.config import settings

# Module-level httpx client with connection pooling for LanguageTool
_lt_client = httpx.Client(
    limits=httpx.Limits(max_connections=20, max_keepalive_connections=10),
    timeout=httpx.Timeout(settings.lt_timeout),
)
from app.utils import minio_client
from app.utils.openai_client import openai_client
from app.services.prompt_builder import build_system_prompt, build_user_prompt
from app.services.complexity_router import (
    route_paragraph, compute_section_position, CorrectionRoute,
)
from app.services.quality_gates import validate_correction as run_quality_gates
from app.services.engine_router import decide_engines, revert_lt_changes_in_protected_regions
from app.services.protected_regions import regions_to_prompt_text
from app.services.audit_pass import audit_paragraph_with_context, should_run_pass2

logger = logging.getLogger(__name__)


@dataclass
class CorrectionResult:
    """Resultado de corrección de un bloque de texto."""
    original_text: str
    corrected_text: str
    operations: list[dict]
    has_changes: bool
    source: str  # 'languagetool', 'languagetool+chatgpt'


def correct_text_with_languagetool(
    text: str,
    language: str = "es",
    disabled_rules: list[str] | None = None,
) -> CorrectionResult:
    """
    Corrige texto con LanguageTool servidor local.
    Usa la API HTTP directa (más eficiente que la librería Python).
    """
    if not text or not text.strip():
        return CorrectionResult(
            original_text=text,
            corrected_text=text,
            operations=[],
            has_changes=False,
            source="languagetool",
        )

    # Llamar a LanguageTool API con connection pooling y retry
    url = f"{settings.languagetool_url}/v2/check"
    payload = {
        "text": text,
        "language": language,
    }
    if disabled_rules:
        payload["disabledRules"] = ",".join(disabled_rules)

    @retry(
        stop=stop_after_attempt(settings.lt_max_retries + 1),
        wait=wait_fixed(1),
        retry=retry_if_exception_type((httpx.TimeoutException, httpx.ConnectError)),
        before_sleep=lambda rs: logger.warning(
            f"LanguageTool retry #{rs.attempt_number} after {type(rs.outcome.exception()).__name__}"
        ),
        reraise=True,
    )
    def _call_lt():
        resp = _lt_client.post(url, data=payload)
        resp.raise_for_status()
        return resp.json()

    try:
        result = _call_lt()
    except (httpx.HTTPError, Exception) as e:
        logger.error(f"Error llamando a LanguageTool: {e}")
        return CorrectionResult(
            original_text=text,
            corrected_text=text,
            operations=[],
            has_changes=False,
            source="languagetool",
        )

    matches = result.get("matches", [])
    if not matches:
        return CorrectionResult(
            original_text=text,
            corrected_text=text,
            operations=[],
            has_changes=False,
            source="languagetool",
        )

    # Aplicar correcciones de atrás hacia adelante (para no alterar offsets)
    operations = []
    corrected = text

    # Ordenar matches por offset descendente para aplicar de atrás hacia adelante
    sorted_matches = sorted(matches, key=lambda m: m["offset"], reverse=True)

    for match in sorted_matches:
        replacements = match.get("replacements", [])
        if not replacements:
            continue

        offset = match["offset"]
        length = match["length"]
        original_fragment = corrected[offset:offset + length]
        replacement = replacements[0]["value"]  # Tomar la primera sugerencia

        # Aplicar reemplazo
        corrected = corrected[:offset] + replacement + corrected[offset + length:]

        operations.append({
            "offset": offset,
            "length": length,
            "original": original_fragment,
            "replacement": replacement,
            "rule_id": match.get("rule", {}).get("id", ""),
            "category": match.get("rule", {}).get("category", {}).get("id", ""),
            "message": match.get("message", ""),
        })

    # Revertir orden para que operations estén en orden natural
    operations.reverse()

    has_changes = corrected != text

    return CorrectionResult(
        original_text=text,
        corrected_text=corrected,
        operations=operations,
        has_changes=has_changes,
        source="languagetool",
    )


def correct_page_blocks_sync(
    doc_id: str,
    page_no: int,
    blocks: list[dict],
    config: dict,
) -> list[dict]:
    """
    Corrige todos los bloques de texto de una página (extracción PDF).
    Se mantiene para registrar patches en BD y para futuras Rutas 2/3.
    Para Ruta 1 DOCX-first, la corrección real se hace en correct_docx_sync.
    """
    language = config.get("language", "es")
    disabled_rules = config.get("lt_disabled_rules", [])

    patches = []

    for block in blocks:
        if block.get("type") != "text":
            continue

        text = block.get("text", "").strip()
        if not text or len(text) < 3:
            continue

        result = correct_text_with_languagetool(
            text=text,
            language=language,
            disabled_rules=disabled_rules,
        )

        if result.has_changes:
            patch_data = {
                "block_no": block["block_no"],
                "original_text": result.original_text,
                "corrected_text": result.corrected_text,
                "operations": result.operations,
                "source": result.source,
                "bbox": block["bbox"],
            }
            patches.append(patch_data)

            logger.info(
                f"Página {page_no}, bloque {block['block_no']}: "
                f"{len(result.operations)} correcciones LT"
            )

    # Guardar parches como JSON en MinIO
    if patches:
        patch_key = f"pages/{doc_id}/patch/{page_no}_v1.json"
        patch_bytes = json.dumps(patches, ensure_ascii=False, indent=2).encode("utf-8")
        minio_client.upload_file(patch_key, patch_bytes, content_type="application/json")

    logger.info(f"Página {page_no}: {len(patches)} bloques con correcciones")
    return patches


# =====================================================================
# CORRECCIÓN DIRECTA DE DOCX (Ruta 1 — evita fragmentación del PDF)
# =====================================================================


def _correct_single_paragraph(
    idx: int,
    para_text: str,
    location: str,
    language: str,
    disabled_rules: list[str],
    profile: dict | None,
    system_prompt: str | None,
    max_expansion: float,
    sections: list[dict],
    para_classifications: dict[int, dict],
    context_prev: str | dict | None,
    context_window: list[str] | None = None,
    precomputed_lt: dict | None = None,
    next_paragraph_type: str | None = None,
    table_context: dict | None = None,
    has_page_break: bool = False,
    audit_log_collector: list | None = None,
) -> tuple[dict | None, dict | None, str, str]:
    """
    Corrige un párrafo individual.
    Reutilizado por la ruta secuencial, los batch tasks y el boundary check.

    Args:
        context_prev: Último párrafo corregido — puede ser str (legacy) o dict enriquecido
            {text, type, ends_abruptly, location_type} para mayor contexto estructural.
        context_window: Últimos N párrafos corregidos (seed para fallback MVP1).
        precomputed_lt: Resultado pre-computado de LT (evita llamada HTTP). Tiene las
            claves: corrected_text, operations, has_changes. Si es None se llama a LT.
        next_paragraph_type: Tipo del párrafo siguiente (lookahead para no alterar flujo).
        table_context: Contexto de tabla {header_row, num_cols, col_index} para celdas.

    Returns:
        Tupla (patch_data | None, usage_record | None, final_text, route_taken).
        patch_data es None si no hay cambios respecto al texto original.
        route_taken siempre se retorna para acumulación de estadísticas.
    """
    text = para_text.strip()

    # === Lote 4: Clasificación — necesaria antes de engines para las reglas de decisión ===
    classification = para_classifications.get(idx, {})
    paragraph_type = classification.get("paragraph_type")

    # Fallback: inferir tipo desde location string si no hay clasificación
    if not paragraph_type:
        if location.startswith("header:"):
            paragraph_type = "encabezado"
        elif location.startswith("footer:"):
            paragraph_type = "footer"
        elif location.startswith("table:"):
            paragraph_type = "celda_tabla"

    # === Sprint 3: Engine Router — detectar regiones protegidas y reglas de colisión ===
    engine_decision = decide_engines(
        text=text,
        paragraph_type=paragraph_type,
        profile=profile,
        base_disabled_rules=disabled_rules,
    )
    protected_regions = engine_decision.protected_regions
    effective_disabled_rules = engine_decision.lt_disabled_rules
    reverted_lt: list[dict] = []

    # === PASO 1: LanguageTool (ortografía y gramática) ===
    if precomputed_lt is not None:
        # Usar resultado pre-computado (Pass 1 paralelo) — los disabled_rules ya se aplicaron
        post_lt_text = precomputed_lt["corrected_text"]
        lt_operations = precomputed_lt["operations"]
    else:
        lt_result = correct_text_with_languagetool(
            text=text,
            language=language,
            disabled_rules=effective_disabled_rules,
        )
        post_lt_text = lt_result.corrected_text
        lt_operations = lt_result.operations

    # Revertir cambios LT que afecten regiones protegidas
    if protected_regions and post_lt_text != text:
        post_lt_text, reverted_lt = revert_lt_changes_in_protected_regions(
            original_text=text,
            lt_corrected=post_lt_text,
            protected_regions=protected_regions,
        )
        if reverted_lt:
            logger.info(
                f"Párrafo {idx}: {len(reverted_lt)} correcciones LT revertidas "
                f"(región protegida)"
            )

    source = "languagetool"

    if lt_operations:
        logger.info(f"Párrafo {idx}: LanguageTool → {len(lt_operations)} correcciones")

    section_position, is_section_transition, current_section = \
        compute_section_position(idx, sections)

    section_summary = None
    active_terms = None
    if current_section:
        section_summary = current_section.get("summary_text") or current_section.get("topic")
        active_terms = current_section.get("active_terms", [])

    route_decision = route_paragraph(
        text=post_lt_text,
        paragraph_type=paragraph_type,
        lt_matches_count=len(lt_operations),
        profile=profile,
        is_section_transition=is_section_transition,
        section_position=section_position,
    )
    route_taken = route_decision.route.value

    # === PASO 2: LLM (estilo) — según ruta ===
    max_length = int(len(post_lt_text) * max_expansion)
    llm_changes = []
    llm_confidence = None
    llm_rewrite_ratio = None
    model_used_actual = "languagetool"
    usage_record = None

    if route_decision.route == CorrectionRoute.SKIP:
        final_text = post_lt_text

    elif profile and system_prompt:
        # MVP2: Prompt parametrizado con perfil + contexto jerárquico enriquecido
        protected_regions_text = regions_to_prompt_text(protected_regions) if protected_regions else None
        user_prompt = build_user_prompt(
            text=post_lt_text,
            profile=profile,
            context_prev=context_prev,
            paragraph_index=idx,
            section_summary=section_summary,
            active_terms=active_terms,
            paragraph_type=paragraph_type,
            next_paragraph_type=next_paragraph_type,
            table_context=table_context,
            has_page_break=has_page_break,
            protected_regions_text=protected_regions_text,
        )

        if route_decision.route == CorrectionRoute.EDITORIAL:
            model_override = settings.openai_editorial_model
            max_tokens_for_route = settings.openai_editorial_max_tokens
        else:
            model_override = settings.openai_cheap_model
            max_tokens_for_route = settings.openai_cheap_max_tokens
        model_used_actual = model_override

        def _p1_audit_cb(raw: dict) -> None:
            if audit_log_collector is not None:
                audit_log_collector.append({
                    **raw,
                    "paragraph_index": idx,
                    "location": location,
                    "pass_number": 1,
                    "call_purpose": "mechanical_correction",
                })

        llm_response, llm_usage = openai_client.correct_with_profile(
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            max_length=max_length,
            model_override=model_override,
            max_tokens_override=max_tokens_for_route,
            on_audit_log=_p1_audit_cb,
        )
        if llm_usage["total_tokens"] > 0:
            _cost = (llm_usage["prompt_tokens"] / 1e6 * settings.openai_pricing_input) + \
                    (llm_usage["completion_tokens"] / 1e6 * settings.openai_pricing_output)
            usage_record = {
                "paragraph_index": idx,
                "location": location,
                "call_type": f"correction_{route_taken}",
                "model_used": model_override,
                **llm_usage,
                "cost_usd": round(_cost, 8),
            }

        if llm_response is None:
            # LLM no respondió (timeout, error de modelo, API key inválida, etc.)
            # Visibilidad explícita en logs — antes era un fallback silencioso.
            logger.warning(
                f"Párrafo {idx}: LLM no respondió (modelo={model_override}, ruta={route_taken}). "
                f"Aplicando solo LanguageTool como fallback."
            )
            final_text = post_lt_text
        elif llm_response.get("action") == "correct":
            corrected = llm_response.get("corrected_text", "")
            if corrected and corrected != post_lt_text:
                final_text = corrected
                source = "languagetool+chatgpt"
                llm_changes = llm_response.get("changes", [])
                llm_confidence = llm_response.get("confidence")
                llm_rewrite_ratio = llm_response.get("rewrite_ratio")
                logger.info(
                    f"Párrafo {idx}: LLM ({route_taken}) → "
                    f"{len(llm_changes)} cambios [{route_decision.reason}]"
                )
            else:
                final_text = post_lt_text
        else:
            # LLM respondió pero decidió skip/flag — no es un error
            final_text = post_lt_text
    else:
        # MVP1 fallback: prompt genérico
        chatgpt_text, mvp1_usage = openai_client.correct_text_style(
            original_text=post_lt_text,
            context_blocks=context_window or [],
            max_length_ratio=max_expansion,
        )
        if mvp1_usage["total_tokens"] > 0:
            _cost = (mvp1_usage["prompt_tokens"] / 1e6 * settings.openai_pricing_input) + \
                    (mvp1_usage["completion_tokens"] / 1e6 * settings.openai_pricing_output)
            usage_record = {
                "paragraph_index": idx,
                "location": location,
                "call_type": "correction_mvp1",
                "model_used": settings.openai_model,
                **mvp1_usage,
                "cost_usd": round(_cost, 8),
            }
        model_used_actual = settings.openai_model
        if chatgpt_text is not None and chatgpt_text != post_lt_text:
            final_text = chatgpt_text
            source = "languagetool+chatgpt"
            logger.info(f"Párrafo {idx}: ChatGPT → estilo mejorado")
        else:
            final_text = post_lt_text

    # === Lote 5: Quality Gates ===
    review_status = "auto_accepted"
    review_reason = None
    gate_results_data = []

    if final_text != text and route_decision.route != CorrectionRoute.SKIP:
        gates = run_quality_gates(
            original=text,
            corrected=final_text,
            profile=profile,
            protected_terms=profile.get("protected_terms", []) if profile else [],
            paragraph_type=paragraph_type,
        )
        gate_results_data = [g.to_dict() for g in gates]
        failed_critical = [g for g in gates if not g.passed and g.critical]
        failed_non_critical = [g for g in gates if not g.passed and not g.critical]

        if failed_critical:
            reasons = "; ".join(g.message for g in failed_critical)
            logger.warning(
                f"Párrafo {idx}: gates críticos fallaron → descartando corrección LLM: {reasons}"
            )
            final_text = post_lt_text
            source = "languagetool" if lt_operations else source
            llm_changes = []
            review_status = "gate_rejected"
            review_reason = reasons
        elif failed_non_critical:
            reasons = "; ".join(g.message for g in failed_non_critical)
            review_status = "manual_review"
            review_reason = reasons
            logger.info(
                f"Párrafo {idx}: gates no-críticos fallaron → manual_review: {reasons}"
            )

    # Construir patch_data solo si hay cambios
    patch_data = None
    if final_text != text:
        patch_data = {
            "paragraph_index": idx,
            "location": location,
            "original_text": text,
            "corrected_text": final_text,
            "lt_operations": lt_operations,
            "source": source,
            "changes": llm_changes,
            "confidence": llm_confidence,
            "rewrite_ratio": llm_rewrite_ratio,
            "model_used": model_used_actual if "chatgpt" in source else "languagetool",
            "route_taken": route_taken,
            "review_status": review_status,
            "review_reason": review_reason,
            "gate_results": gate_results_data,
            # Sprint 3: Audit trail dual-engine
            "lt_corrections_json": lt_operations if lt_operations else None,
            "llm_change_log_json": llm_changes if llm_changes else None,
            "reverted_lt_changes_json": reverted_lt if reverted_lt else None,
            "protected_regions_snapshot": [r.to_dict() for r in protected_regions] if protected_regions else None,
        }

    return patch_data, usage_record, final_text, route_taken


def correct_docx_sync(
    doc_id: str,
    docx_uri: str,
    config: dict,
    profile: dict | None = None,
    analysis_data: dict | None = None,
    on_progress: Callable[[int, int], None] | None = None,
    docx_bytes_cached: bytes | None = None,
    global_context: dict | None = None,
) -> tuple[list[dict], list[dict], list[tuple[str, str, bool]], list[dict]]:
    """
    Corrige un DOCX directamente, párrafo por párrafo.
    Plan v4: Pipeline de Doble Pasada — Pasada 1 (mecánica) + Pasada 2 (auditoría contextual).

    Args:
        analysis_data: Resultado de Etapa C con sections, paragraph_classifications.
        global_context: ADN editorial del documento (de Etapa C.6). Si None, no hay Pasada 2.

    Retorna tupla (patches, usage_records, all_paragraphs, audit_log_entries).
    audit_log_entries: lista de payloads RAW de cada llamada LLM para persistir en llm_audit_log.
    """
    language = config.get("language", "es")
    disabled_rules = config.get("lt_disabled_rules", [])

    # Si hay perfil, usar sus disabled_rules también
    if profile and profile.get("lt_disabled_rules"):
        disabled_rules = list(set(disabled_rules + profile["lt_disabled_rules"]))

    # Descargar DOCX (con cache si disponible)
    if docx_bytes_cached is not None:
        docx_bytes = docx_bytes_cached
    else:
        docx_bytes = minio_client.download_file(docx_uri)
    tmpfile = tempfile.mktemp(suffix=".docx")
    with open(tmpfile, "wb") as f:
        f.write(docx_bytes)

    doc = DocxDocument(tmpfile)
    Path(tmpfile).unlink(missing_ok=True)

    patches = []
    corrected_context: list[str] = []
    corrected_meta: list[dict] = []   # contexto enriquecido con tipo y metadata
    usage_records: list[dict] = []
    audit_log_entries: list[dict] = []  # Plan v4: payloads RAW para llm_audit_log

    all_paragraphs = _collect_all_paragraphs(doc)
    logger.info(f"Documento {doc_id}: {len(all_paragraphs)} párrafos a corregir")

    # MVP2: Construir system prompt UNA VEZ (cacheable)
    system_prompt = build_system_prompt() if profile else None
    max_expansion = profile.get("max_expansion_ratio", 1.15) if profile else 1.15

    # Lote 4: Preparar datos del análisis para contexto jerárquico
    sections = (analysis_data or {}).get("sections", [])
    para_classifications = {}
    for pc in (analysis_data or {}).get("paragraph_classifications", []):
        para_classifications[pc["paragraph_index"]] = pc

    # Construir mapa de contexto de tabla para corrección table-aware
    table_context_map = _build_table_context_map(doc)

    # Contadores para logging final
    route_counts = {"skip": 0, "cheap": 0, "editorial": 0}
    gate_stats = {"passed": 0, "discarded": 0, "flagged": 0}

    # Fase 2: Pass 1 LT en paralelo (cuando hay más de 1 worker configurado).
    # LT es stateless → sin dependencias inter-párrafo → seguro paralelizar.
    lt_precomputed: list[dict] = []
    lt_workers = settings.parallel_correction_lt_workers
    if lt_workers > 1 and len(all_paragraphs) > 1:
        lt_precomputed = correct_all_paragraphs_lt_sync(
            all_paragraphs=all_paragraphs,
            language=language,
            disabled_rules=disabled_rules,
            max_workers=lt_workers,
        )

    for idx, (para_text, location, para_has_pb) in enumerate(all_paragraphs):
        text = para_text.strip()
        if not text or len(text) < 3:
            continue

        # Contexto previo enriquecido con metadata estructural
        context_prev: dict | None = None
        if corrected_meta:
            last = corrected_meta[-1]
            last_text = last["text"]
            context_prev = {
                "text": last_text,
                "type": last.get("type", ""),
                "ends_abruptly": bool(last_text) and last_text[-1] not in '.?!:)"»\n',
                "location_type": last.get("location", "body:").split(":")[0],
            }

        # Lookahead: tipo del párrafo siguiente
        next_para_type: str | None = None
        if idx + 1 < len(all_paragraphs):
            _, next_loc, _ = all_paragraphs[idx + 1]
            next_cls = para_classifications.get(idx + 1, {})
            next_para_type = next_cls.get("paragraph_type")
            if not next_para_type:
                if next_loc.startswith("header:"):
                    next_para_type = "encabezado"
                elif next_loc.startswith("footer:"):
                    next_para_type = "footer"
                elif next_loc.startswith("table:"):
                    next_para_type = "celda_tabla"

        # Contexto de tabla para celdas
        table_ctx: dict | None = None
        if location.startswith("table:"):
            parts_loc = location.split(":")
            if len(parts_loc) >= 4:
                t_key = f"table:{parts_loc[1]}"
                col_index = int(parts_loc[3]) if len(parts_loc) > 3 else 0
                if t_key in table_context_map:
                    table_ctx = {**table_context_map[t_key], "col_index": col_index}

        precomputed_lt = lt_precomputed[idx] if lt_precomputed else None
        patch_data, usage_record, final_text, route_taken = _correct_single_paragraph(
            idx=idx,
            para_text=para_text,
            location=location,
            language=language,
            disabled_rules=disabled_rules,
            profile=profile,
            system_prompt=system_prompt,
            max_expansion=max_expansion,
            sections=sections,
            para_classifications=para_classifications,
            context_prev=context_prev,
            context_window=corrected_context[-3:],
            precomputed_lt=precomputed_lt,
            next_paragraph_type=next_para_type,
            table_context=table_ctx,
            has_page_break=para_has_pb,
            audit_log_collector=audit_log_entries,
        )

        # ── Plan v4: PASADA 2 — Auditoría Contextual ────────────────────
        classification = para_classifications.get(idx, {})
        paragraph_type = classification.get("paragraph_type")
        p1_rewrite_ratio = patch_data.get("rewrite_ratio") if patch_data else None
        p1_has_changes = patch_data is not None
        intervention_level = (profile or {}).get("intervention_level")

        if global_context and should_run_pass2(
            route_taken=route_taken,
            pass1_rewrite_ratio=p1_rewrite_ratio,
            intervention_level=intervention_level,
            pass1_has_changes=p1_has_changes,
        ):
            corrected_pass1 = final_text  # resultado de Pasada 1

            def _p2_audit_cb(raw: dict) -> None:
                audit_log_entries.append({
                    **raw,
                    "paragraph_index": idx,
                    "location": location,
                    "pass_number": 2,
                    "call_purpose": "contextual_audit",
                })

            audit_result, p2_usage = audit_paragraph_with_context(
                original_text=text,
                corrected_pass1=corrected_pass1,
                global_context=global_context,
                context_window=corrected_context[-3:],
                paragraph_type=paragraph_type,
                location=location,
                on_audit_log=_p2_audit_cb,
            )

            if audit_result and audit_result.get("final_text"):
                final_text = audit_result["final_text"]
                # Calcular costo Pasada 2
                if p2_usage.get("total_tokens", 0) > 0:
                    _p2_cost = (
                        p2_usage["prompt_tokens"] / 1e6 * settings.openai_pricing_input
                        + p2_usage["completion_tokens"] / 1e6 * settings.openai_pricing_output
                    )
                    usage_records.append({
                        "paragraph_index": idx,
                        "location": location,
                        "call_type": "audit_pass2",
                        "model_used": settings.openai_editorial_model,
                        **p2_usage,
                        "cost_usd": round(_p2_cost, 8),
                    })
                # Actualizar o crear patch_data con resultado de Pasada 2
                if patch_data is None and final_text != text:
                    patch_data = {
                        "paragraph_index": idx,
                        "location": location,
                        "original_text": text,
                        "corrected_text": final_text,
                        "lt_operations": [],
                        "source": "chatgpt",
                        "changes": audit_result.get("style_improvements", []),
                        "confidence": audit_result.get("confidence"),
                        "rewrite_ratio": None,
                        "model_used": settings.openai_editorial_model,
                        "route_taken": route_taken,
                        "review_status": "auto_accepted",
                        "review_reason": None,
                        "gate_results": [],
                        "lt_corrections_json": None,
                        "llm_change_log_json": None,
                        "reverted_lt_changes_json": None,
                        "protected_regions_snapshot": None,
                    }
                if patch_data is not None:
                    patch_data["corrected_pass1_text"] = corrected_pass1
                    patch_data["pass2_audit_json"] = {
                        "reverted_destructions": audit_result.get("reverted_destructions", []),
                        "style_improvements": audit_result.get("style_improvements", []),
                        "confidence": audit_result.get("confidence"),
                        "pass1_quality": audit_result.get("pass1_quality"),
                    }
                    # Solo actualizar corrected_text si Pasada 2 produjo resultado diferente
                    if final_text != corrected_pass1:
                        patch_data["corrected_text"] = final_text
                        patch_data["source"] = (
                            patch_data.get("source", "languagetool") + "+audit"
                        ).lstrip("+")

        corrected_context.append(final_text)
        # Guardar metadata para el siguiente contexto enriquecido
        cls = para_classifications.get(idx, {})
        corrected_meta.append({
            "text": final_text,
            "type": cls.get("paragraph_type", ""),
            "location": location,
        })
        route_counts[route_taken] += 1

        if patch_data:
            patches.append(patch_data)
            # Gate stats solo para rutas donde se ejecutaron quality gates (no SKIP)
            if route_taken != "skip":
                rs = patch_data["review_status"]
                if rs == "gate_rejected":
                    gate_stats["discarded"] += 1
                elif rs == "manual_review":
                    gate_stats["flagged"] += 1
                else:
                    gate_stats["passed"] += 1

        if usage_record:
            usage_records.append(usage_record)

        if on_progress:
            on_progress(idx + 1, len(all_paragraphs))

    # Guardar parches en MinIO
    if patches:
        patch_key = f"docx/{doc_id}/patches_docx.json"
        patch_bytes = json.dumps(patches, ensure_ascii=False, indent=2).encode("utf-8")
        minio_client.upload_file(patch_key, patch_bytes, content_type="application/json")

    total_tokens = sum(r["total_tokens"] for r in usage_records)
    total_cost = sum(r["cost_usd"] for r in usage_records)
    logger.info(
        f"Documento {doc_id}: {len(patches)} párrafos corregidos "
        f"({sum(1 for p in patches if 'chatgpt' in p['source'])} con GPT), "
        f"tokens: {total_tokens}, costo: ${total_cost:.6f}, "
        f"llamadas: {len(usage_records)}, "
        f"rutas: skip={route_counts['skip']} cheap={route_counts['cheap']} editorial={route_counts['editorial']}, "
        f"gates: ok={gate_stats['passed']} descartados={gate_stats['discarded']} revisión={gate_stats['flagged']}"
    )
    return patches, usage_records, all_paragraphs, audit_log_entries


def _build_table_context_map(doc: DocxDocument) -> dict[str, dict]:
    """
    Construye un mapa de contexto por tabla para corrección table-aware.
    Clave: 'table:T_IDX', valor: {num_cols, header_row, num_rows}.
    """
    result: dict[str, dict] = {}
    for t_idx, table in enumerate(doc.tables):
        rows = table.rows
        if not rows:
            continue
        header_row = [cell.text.strip() for cell in rows[0].cells]
        result[f"table:{t_idx}"] = {
            "num_cols": len(rows[0].cells),
            "num_rows": len(rows),
            "header_row": header_row,
        }
    return result


_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _has_internal_page_break(para) -> bool:
    """Devuelve True si el párrafo contiene un <w:br w:type="page"/> interno."""
    for br in para._p.findall(f".//{{{_WORD_NS}}}br"):
        if br.get(f"{{{_WORD_NS}}}type") == "page":
            return True
    return False


def _collect_all_paragraphs(doc: DocxDocument) -> list[tuple[str, str, bool]]:
    """
    Recolecta todos los párrafos del documento con su ubicación y flag de salto de página.
    Retorna lista de (texto, ubicación, has_page_break) donde ubicación es:
    - 'body:N' para párrafos del cuerpo
    - 'table:T:R:C:P' para párrafos en tablas
    - 'header:S:P' / 'footer:S:P' para encabezados/pies
    """
    paragraphs: list[tuple[str, str, bool]] = []

    # Cuerpo principal
    for i, para in enumerate(doc.paragraphs):
        paragraphs.append((para.text, f"body:{i}", _has_internal_page_break(para)))

    # Tablas
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    paragraphs.append(
                        (para.text, f"table:{t_idx}:{r_idx}:{c_idx}:{p_idx}", False)
                    )

    # Headers y footers
    for s_idx, section in enumerate(doc.sections):
        for hf_type, hf in [("header", section.header), ("footer", section.footer)]:
            if hf is None:
                continue
            for p_idx, para in enumerate(hf.paragraphs):
                paragraphs.append((para.text, f"{hf_type}:{s_idx}:{p_idx}", False))

    return paragraphs


def save_paragraph_locations_sync(
    doc_id: str,
    all_paragraphs: list[tuple[str, str, bool]],
    para_classifications: dict[int, dict],
    total_pages: int,
    db,
) -> int:
    """
    Persiste paragraph_locations para un documento.
    Usa heurística lineal para page_start (mejora futura: correlación con bloques PDF).
    Retorna el número de registros creados.
    """
    from app.models.paragraph_location import ParagraphLocation
    from sqlalchemy import delete as sa_delete

    total_paras = max(len(all_paragraphs), 1)

    # Limpiar registros previos del documento
    db.execute(sa_delete(ParagraphLocation).where(ParagraphLocation.doc_id == doc_id))

    records = []
    for idx, (_, location, has_pb) in enumerate(all_paragraphs):
        cls = para_classifications.get(idx, {})
        para_type = cls.get("paragraph_type")

        # Heurística lineal: asume distribución uniforme de párrafos en páginas
        est_page = max(1, min(int(idx / total_paras * total_pages) + 1, total_pages))

        records.append(ParagraphLocation(
            doc_id=doc_id,
            paragraph_index=idx,
            location=location,
            page_start=est_page,
            page_end=est_page,
            has_internal_page_break=has_pb,
            paragraph_type=para_type,
            position_in_page="middle",
        ))

    db.bulk_save_objects(records)
    db.commit()
    logger.info(f"Documento {doc_id}: {len(records)} paragraph_locations guardadas")
    return len(records)


def compute_batch_boundaries(
    sections: list[dict],
    all_paragraphs: list[tuple],
    target_batch_size: int,
) -> list[tuple[int, int]]:
    """
    Divide párrafos en lotes alineados a límites de sección.
    Garantiza que ninguna sección quede dividida entre dos lotes.

    Args:
        sections: Lista de secciones de Stage C, cada una con start_paragraph/end_paragraph.
        all_paragraphs: Lista completa de (texto, ubicación).
        target_batch_size: Tamaño objetivo de cada lote en párrafos.

    Returns:
        Lista de (start_idx, end_idx) inclusivas (0-indexed).
        Si no hay secciones o el documento es pequeño, retorna un solo lote.
    """
    total = len(all_paragraphs)
    if total == 0:
        return []
    if not sections or total <= target_batch_size:
        return [(0, total - 1)]

    # Puntos de corte válidos: finales de sección dentro del rango
    valid_cut_points = sorted({
        sec["end_paragraph"]
        for sec in sections
        if "end_paragraph" in sec and 0 <= sec["end_paragraph"] < total
    })

    # Asegurar que el último párrafo siempre sea un punto de corte
    if not valid_cut_points or valid_cut_points[-1] < total - 1:
        valid_cut_points.append(total - 1)

    batches: list[tuple[int, int]] = []
    current_start = 0
    accumulated = 0

    for cut_point in valid_cut_points:
        accumulated += cut_point - current_start + 1
        if accumulated >= target_batch_size or cut_point == total - 1:
            batches.append((current_start, cut_point))
            current_start = cut_point + 1
            accumulated = 0

    return batches


def correct_batch_with_llm_sync(
    batch_index: int,
    start_para: int,
    end_para: int,
    lt_results: list[dict],
    all_paragraphs: list[tuple[str, str, bool]],
    language: str,
    disabled_rules: list[str],
    profile: dict | None,
    system_prompt: str | None,
    max_expansion: float,
    sections: list[dict],
    para_classifications: dict[int, dict],
    context_seed: str | None,
    global_context: dict | None = None,
) -> tuple[list[dict], list[dict], str, list[dict]]:
    """
    Pass 1 LLM + Pass 2 auditoría para el rango [start_para..end_para] (inclusive).
    Usa resultados LT pre-computados. Secuencial dentro del batch (preserva contexto LLM).

    Args:
        lt_results: Lista de resultados LT para todos los párrafos (indexada globalmente).
        context_seed: Texto del último párrafo corregido del batch anterior (aprox.).
        global_context: ADN editorial del documento (Plan v4). Si None, no hay Pasada 2.

    Returns:
        (patches, usage_records, last_corrected_text, audit_log_entries)
        last_corrected_text: último párrafo no vacío corregido — seed para el siguiente batch.
        audit_log_entries: payloads RAW de llamadas LLM para llm_audit_log.
    """
    patches: list[dict] = []
    usage_records: list[dict] = []
    audit_log_entries: list[dict] = []
    corrected_context: list[str] = [context_seed] if context_seed else []
    last_corrected_text: str = context_seed or ""

    intervention_level = (profile or {}).get("intervention_level")

    for idx in range(start_para, min(end_para + 1, len(all_paragraphs))):
        para_text, location, para_has_pb = all_paragraphs[idx]
        text = para_text.strip()
        if not text or len(text) < 3:
            continue

        precomputed_lt = lt_results[idx] if idx < len(lt_results) else None
        context_prev = corrected_context[-1] if corrected_context else None

        patch_data, usage_record, final_text, _ = _correct_single_paragraph(
            idx=idx,
            para_text=para_text,
            location=location,
            language=language,
            disabled_rules=disabled_rules,
            profile=profile,
            system_prompt=system_prompt,
            max_expansion=max_expansion,
            sections=sections,
            para_classifications=para_classifications,
            context_prev=context_prev,
            context_window=corrected_context[-3:],
            precomputed_lt=precomputed_lt,
            has_page_break=para_has_pb,
            audit_log_collector=audit_log_entries,
        )

        # ── Plan v4: PASADA 2 — Auditoría Contextual ────────────────────
        route_taken = patch_data.get("route_taken") if patch_data else "skip"
        p1_rewrite_ratio = patch_data.get("rewrite_ratio") if patch_data else None
        p1_has_changes = patch_data is not None
        classification = para_classifications.get(idx, {})
        paragraph_type = classification.get("paragraph_type")

        if global_context and should_run_pass2(
            route_taken=route_taken,
            pass1_rewrite_ratio=p1_rewrite_ratio,
            intervention_level=intervention_level,
            pass1_has_changes=p1_has_changes,
        ):
            corrected_pass1 = final_text

            def _p2_audit_cb(raw: dict, _idx=idx, _loc=location) -> None:
                audit_log_entries.append({
                    **raw,
                    "paragraph_index": _idx,
                    "location": _loc,
                    "pass_number": 2,
                    "call_purpose": "contextual_audit",
                })

            audit_result, p2_usage = audit_paragraph_with_context(
                original_text=text,
                corrected_pass1=corrected_pass1,
                global_context=global_context,
                context_window=corrected_context[-3:],
                paragraph_type=paragraph_type,
                location=location,
                on_audit_log=_p2_audit_cb,
            )

            if audit_result and audit_result.get("final_text"):
                final_text = audit_result["final_text"]
                if p2_usage.get("total_tokens", 0) > 0:
                    _p2_cost = (
                        p2_usage["prompt_tokens"] / 1e6 * settings.openai_pricing_input
                        + p2_usage["completion_tokens"] / 1e6 * settings.openai_pricing_output
                    )
                    usage_records.append({
                        "paragraph_index": idx,
                        "location": location,
                        "call_type": "audit_pass2",
                        "model_used": settings.openai_editorial_model,
                        **p2_usage,
                        "cost_usd": round(_p2_cost, 8),
                    })
                if patch_data is None and final_text != text:
                    patch_data = {
                        "paragraph_index": idx,
                        "location": location,
                        "original_text": text,
                        "corrected_text": final_text,
                        "lt_operations": [],
                        "source": "chatgpt+audit",
                        "changes": audit_result.get("style_improvements", []),
                        "confidence": audit_result.get("confidence"),
                        "rewrite_ratio": None,
                        "model_used": settings.openai_editorial_model,
                        "route_taken": route_taken,
                        "review_status": "auto_accepted",
                        "review_reason": None,
                        "gate_results": [],
                        "lt_corrections_json": None,
                        "llm_change_log_json": None,
                        "reverted_lt_changes_json": None,
                        "protected_regions_snapshot": None,
                    }
                if patch_data is not None:
                    patch_data["corrected_pass1_text"] = corrected_pass1
                    patch_data["pass2_audit_json"] = {
                        "reverted_destructions": audit_result.get("reverted_destructions", []),
                        "style_improvements": audit_result.get("style_improvements", []),
                        "confidence": audit_result.get("confidence"),
                        "pass1_quality": audit_result.get("pass1_quality"),
                    }
                    if final_text != corrected_pass1:
                        patch_data["corrected_text"] = final_text
                        patch_data["source"] = (
                            patch_data.get("source", "languagetool") + "+audit"
                        ).lstrip("+")

        corrected_context.append(final_text)
        last_corrected_text = final_text

        if patch_data:
            patches.append(patch_data)
        if usage_record:
            usage_records.append(usage_record)

    p2_count = sum(1 for e in audit_log_entries if e.get("pass_number") == 2)
    logger.info(
        f"Batch {batch_index} [{start_para}-{end_para}]: {len(patches)} parches, "
        f"{p2_count} auditorías P2"
    )
    return patches, usage_records, last_corrected_text, audit_log_entries


def check_batch_boundaries(
    batch_results: dict[int, dict],
    batch_boundaries: list[tuple[int, int]],
    lt_results: list[dict],
    all_paragraphs: list[tuple[str, str, bool]],
    language: str,
    disabled_rules: list[str],
    profile: dict | None,
    system_prompt: str | None,
    max_expansion: float,
    sections: list[dict],
    para_classifications: dict[int, dict],
    all_patches: list[dict],
) -> list[dict]:
    """
    Re-corrige el primer párrafo de cada batch (batch > 0) usando el seed real
    (último párrafo corregido del batch anterior) en vez del seed aproximado post-LT.

    Coste: N_batches-1 llamadas LLM (≤7 para 8 lotes, ~35s extra).
    Solo re-corre párrafos donde el seed real difiere del seed aproximado post-LT.
    """
    if len(batch_boundaries) <= 1:
        return all_patches

    # Construir índice paragraph_index → posición en all_patches para acceso O(1)
    patch_index: dict[int, int] = {
        p["paragraph_index"]: i for i, p in enumerate(all_patches)
    }

    updated_patches = list(all_patches)  # copia shallow para mutaciones seguras
    recorrected = 0

    for b_idx in range(1, len(batch_boundaries)):
        start_para, end_para = batch_boundaries[b_idx]
        prev_batch_data = batch_results.get(b_idx - 1)
        if not prev_batch_data:
            logger.warning(f"Boundary check: no hay datos del batch {b_idx - 1}, saltando")
            continue

        real_seed = prev_batch_data.get("last_corrected_text") or ""

        # Encontrar el primer párrafo no vacío del batch actual
        first_idx = None
        for para_idx in range(start_para, end_para + 1):
            if para_idx < len(all_paragraphs):
                para_text = all_paragraphs[para_idx][0]
                if para_text.strip() and len(para_text.strip()) >= 3:
                    first_idx = para_idx
                    break

        if first_idx is None:
            logger.debug(f"Boundary check batch {b_idx}: sin párrafos válidos")
            continue

        # Seed aproximado que usó el batch original (post-LT del párrafo anterior)
        approx_seed = ""
        if first_idx > 0 and lt_results and (first_idx - 1) < len(lt_results):
            approx_seed = (lt_results[first_idx - 1].get("corrected_text") or "")[:200]

        real_seed_trimmed = real_seed[:200] if real_seed else ""

        if real_seed_trimmed == approx_seed:
            logger.debug(
                f"Boundary check batch {b_idx}, párr {first_idx}: "
                f"seed real == seed aprox, sin re-corrección necesaria"
            )
            continue

        # Re-corregir con el seed real
        para_text, location, para_has_pb = all_paragraphs[first_idx]
        precomputed_lt = lt_results[first_idx] if lt_results and first_idx < len(lt_results) else None

        try:
            new_patch_data, _usage, _final_text, _route = _correct_single_paragraph(
                idx=first_idx,
                para_text=para_text,
                location=location,
                language=language,
                disabled_rules=disabled_rules,
                profile=profile,
                system_prompt=system_prompt,
                max_expansion=max_expansion,
                sections=sections,
                para_classifications=para_classifications,
                context_prev=real_seed_trimmed,
                precomputed_lt=precomputed_lt,
                has_page_break=para_has_pb,
            )
        except Exception as e:
            logger.error(
                f"Boundary check batch {b_idx}, párr {first_idx}: error en re-corrección: {e}"
            )
            continue

        existing_pos = patch_index.get(first_idx)

        if new_patch_data is None:
            # Re-corrección no produjo cambios: eliminar patch previo si existía
            if existing_pos is not None:
                old_patch = updated_patches[existing_pos]
                if old_patch.get("corrected_text") != para_text.strip():
                    logger.info(
                        f"Boundary check batch {b_idx}, párr {first_idx}: "
                        f"re-corrección sin cambios → eliminando patch previo"
                    )
                    updated_patches[existing_pos] = None  # marcado para filtrar
                    recorrected += 1
        else:
            if existing_pos is not None:
                old_text = updated_patches[existing_pos].get("corrected_text", "")
                if old_text != new_patch_data["corrected_text"]:
                    updated_patches[existing_pos] = new_patch_data
                    recorrected += 1
                    logger.info(
                        f"Boundary check batch {b_idx}, párr {first_idx}: "
                        f"patch actualizado con seed real"
                    )
            else:
                # El párrafo no tenía patch antes, ahora sí
                updated_patches.append(new_patch_data)
                patch_index[first_idx] = len(updated_patches) - 1
                recorrected += 1
                logger.info(
                    f"Boundary check batch {b_idx}, párr {first_idx}: "
                    f"nuevo patch generado con seed real"
                )

    # Filtrar entradas marcadas como None y re-ordenar
    result = [p for p in updated_patches if p is not None]
    result.sort(key=lambda p: p["paragraph_index"])

    logger.info(
        f"check_batch_boundaries: {len(batch_boundaries)} lotes, "
        f"{recorrected} párrafos de frontera re-corregidos"
    )
    return result


def correct_all_paragraphs_lt_sync(
    all_paragraphs: list[tuple[str, str, bool]],
    language: str,
    disabled_rules: list[str],
    max_workers: int,
) -> list[dict]:
    """
    Pass 1: LanguageTool en paralelo para todos los párrafos.
    LanguageTool es 100% stateless — sin dependencias inter-párrafo.

    Returns list of LT result dicts indexed by paragraph position.
    Each dict: {text, corrected_text, operations, has_changes, skip}.
    Short/empty paragraphs are marked with skip=True.
    """
    results: list[dict | None] = [None] * len(all_paragraphs)

    def _run_lt(idx: int) -> tuple[int, dict]:
        para_text = all_paragraphs[idx][0]
        text = para_text.strip()
        if not text or len(text) < 3:
            return idx, {
                "text": para_text,
                "corrected_text": para_text,
                "operations": [],
                "has_changes": False,
                "skip": True,
            }
        lt_result = correct_text_with_languagetool(
            text=text,
            language=language,
            disabled_rules=disabled_rules,
        )
        return idx, {
            "text": text,
            "corrected_text": lt_result.corrected_text,
            "operations": lt_result.operations,
            "has_changes": lt_result.has_changes,
            "skip": False,
        }

    with ThreadPoolExecutor(max_workers=max_workers) as pool:
        futures = {pool.submit(_run_lt, i): i for i in range(len(all_paragraphs))}
        for future in as_completed(futures):
            idx, result = future.result()
            results[idx] = result

    lt_hits = sum(1 for r in results if r and r["has_changes"])
    logger.info(
        f"Pass 1 LT paralelo: {len(all_paragraphs)} párrafos procesados, "
        f"{lt_hits} con correcciones ({max_workers} workers)"
    )
    return results  # type: ignore[return-value]
