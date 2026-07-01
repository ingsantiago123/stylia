"""
Servicio de Renderizado (Etapa E).
MVP 1: Solo Ruta 1 — DOCX-first.
Aplica correcciones párrafo por párrafo al DOCX, preservando formato de runs.
MVP 2: Genera previews anotados con highlights sobre texto corregido.
"""

import json
import logging
import tempfile
from pathlib import Path

from difflib import SequenceMatcher

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.oxml.ns import qn

from app.utils import minio_client
from app.utils.pdf_utils import convert_docx_to_pdf

logger = logging.getLogger(__name__)


# Colores RGB (0-1) para highlights por categoría editorial
HIGHLIGHT_COLORS = {
    "redundancia": (1.0, 0.65, 0.0),
    "claridad":    (0.3, 0.6, 1.0),
    "registro":    (0.5, 0.4, 1.0),
    "cohesion":    (0.0, 0.75, 0.85),
    "lexico":      (0.0, 0.7, 0.55),
    "estructura":  (0.6, 0.35, 1.0),
    "puntuacion":  (0.95, 0.75, 0.0),
    "ritmo":       (0.9, 0.4, 0.6),
    "muletilla":   (0.9, 0.3, 0.4),
}
DEFAULT_HIGHLIGHT = (0.83, 1.0, 0.0)  # krypton


def _get_patch_metadata(patch: dict) -> dict:
    """Extrae category/severity/explanation de un patch (MVP2 con changes list)."""
    changes = patch.get("changes", [])
    if changes and isinstance(changes, list) and len(changes) > 0:
        first = changes[0]
        return {
            "category": first.get("category", ""),
            "severity": first.get("severity"),
            "explanation": first.get("explanation"),
        }
    return {
        "category": patch.get("category", ""),
        "severity": patch.get("severity"),
        "explanation": patch.get("explanation"),
    }


def _compute_deleted_phrases_in_original(
    original: str, corrected: str, context_words: int = 2, min_len: int = 5
) -> list[str]:
    """
    Diff a nivel de palabras: devuelve frases del texto ORIGINAL que fueron
    eliminadas o reemplazadas, con palabras de contexto para unicidad en búsqueda PDF.
    """
    orig_words = original.split()
    corr_words = corrected.split()
    if not orig_words or not corr_words:
        return []
    matcher = SequenceMatcher(None, orig_words, corr_words, autojunk=False)
    phrases: list[str] = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag not in ("delete", "replace"):
            continue
        deleted = orig_words[i1:i2]
        if not deleted:
            continue
        ctx_before = orig_words[max(0, i1 - context_words):i1]
        ctx_after = orig_words[i2:min(len(orig_words), i2 + context_words)]
        phrase = " ".join(ctx_before + deleted + ctx_after)
        if len(phrase.strip()) >= min_len:
            phrases.append(phrase.strip())
    return phrases


def _compute_changed_phrases_in_corrected(
    original: str, corrected: str, context_words: int = 2, min_len: int = 5
) -> list[str]:
    """
    Word-level diff: returns context-enriched phrases from the corrected text
    that correspond to changed or inserted word regions.
    context_words = unchanged words included on each side to make the phrase unique.
    """
    orig_words = original.split()
    corr_words = corrected.split()
    if not orig_words or not corr_words:
        return []
    matcher = SequenceMatcher(None, orig_words, corr_words, autojunk=False)
    phrases: list[str] = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag not in ("insert", "replace"):
            continue
        changed = corr_words[j1:j2]
        if not changed:
            continue
        ctx_before = corr_words[max(0, j1 - context_words):j1]
        ctx_after = corr_words[j2:min(len(corr_words), j2 + context_words)]
        phrase = " ".join(ctx_before + changed + ctx_after)
        if len(phrase.strip()) >= min_len:
            phrases.append(phrase.strip())
    return phrases


def _generate_annotated_previews(
    doc_id: str,
    corrected_pdf_bytes: bytes,
    all_patches: list[dict],
    render_mode: str = "final",
) -> int:
    """
    Genera previews PNG anotados del PDF corregido.

    Dos capas de anotación por corrección:
    - "paragraph": bbox de la zona del párrafo (outline en frontend, sin highlight baked)
    - "change": bbox exacto de las palabras cambiadas (highlight baked en el PNG)

    render_mode:
    - "candidate": sube a preview_candidate/ y annotations_candidate/
    - "final": sube a preview_corrected/ y annotations/

    Retorna el número total de páginas.
    """
    if render_mode == "candidate":
        preview_prefix = f"pages/{doc_id}/preview_candidate"
        annot_prefix = f"pages/{doc_id}/annotations_candidate"
    else:
        preview_prefix = f"pages/{doc_id}/preview_corrected"
        annot_prefix = f"pages/{doc_id}/annotations"

    pdf_doc = fitz.open(stream=corrected_pdf_bytes, filetype="pdf")
    total_pages = len(pdf_doc)
    page_annotations: dict[int, list] = {p + 1: [] for p in range(total_pages)}
    annotations_found = 0

    total_paragraphs = len(all_patches) or 1

    for patch_idx, patch in enumerate(all_patches):
        orig = patch["original_text"].strip()
        corr = patch["corrected_text"].strip()
        if orig == corr or len(corr) < 3:
            continue

        meta = _get_patch_metadata(patch)
        color = HIGHLIGHT_COLORS.get(meta["category"], DEFAULT_HIGHLIGHT)
        p_ids = patch.get("patch_ids") or ([patch["patch_id"]] if patch.get("patch_id") else [])

        para_idx = patch.get("paragraph_index", patch_idx)
        est_page = min(int(para_idx / total_paragraphs * total_pages), total_pages - 1)
        search_window = 2
        nearby = list(range(max(0, est_page - search_window),
                            min(total_pages, est_page + search_window + 1)))
        remaining = [p for p in range(total_pages) if p not in set(nearby)]

        # Locate the paragraph in the PDF (progressive prefix fallback)
        found_page_idx: int | None = None
        found_quads: list | None = None

        for max_len in [150, 70, 35]:
            search_text = corr[:max_len] if len(corr) > max_len else corr
            if len(search_text) < 4:
                break
            for page_idx in (nearby + remaining):
                quads = pdf_doc[page_idx].search_for(search_text, quads=True)
                if quads:
                    found_page_idx = page_idx
                    found_quads = list(quads)
                    break
            if found_page_idx is not None:
                break

        if found_page_idx is None or not found_quads:
            continue

        annotations_found += 1
        page_no = found_page_idx + 1
        page = pdf_doc[found_page_idx]
        page_rect = page.rect

        ann_base = {
            "patch_ids": p_ids,
            "category": meta["category"],
            "severity": meta["severity"],
            "explanation": meta["explanation"],
            "confidence": patch.get("confidence"),
            "source": patch.get("source", ""),
            "review_status": patch.get("review_status", ""),
            "original_snippet": orig[:100],
            "corrected_snippet": corr[:100],
        }

        # Layer 1: paragraph-level outline (no PDF highlight baked in)
        for quad in found_quads:
            r = quad.rect
            page_annotations[page_no].append({
                **ann_base,
                "annot_type": "paragraph",
                "x_pct": round(r.x0 / page_rect.width * 100, 2),
                "y_pct": round(r.y0 / page_rect.height * 100, 2),
                "w_pct": round((r.x1 - r.x0) / page_rect.width * 100, 2),
                "h_pct": round((r.y1 - r.y0) / page_rect.height * 100, 2),
            })

        # Layer 2: precise changed-word positions (highlight baked into PNG)
        for phrase in _compute_changed_phrases_in_corrected(orig, corr):
            if len(phrase) < 4:
                continue
            change_quads = page.search_for(phrase, quads=True)
            if not change_quads:
                continue
            hl = page.add_highlight_annot(change_quads)
            hl.set_colors(stroke=color)
            hl.set_opacity(0.50)
            hl.update()
            for cquad in change_quads:
                r = cquad.rect
                page_annotations[page_no].append({
                    **ann_base,
                    "annot_type": "change",
                    "x_pct": round(r.x0 / page_rect.width * 100, 2),
                    "y_pct": round(r.y0 / page_rect.height * 100, 2),
                    "w_pct": round((r.x1 - r.x0) / page_rect.width * 100, 2),
                    "h_pct": round((r.y1 - r.y0) / page_rect.height * 100, 2),
                })

    # Render pages as PNG (with baked highlights) and upload with annotation JSON
    for page_idx in range(total_pages):
        page_no = page_idx + 1
        page = pdf_doc[page_idx]
        pix = page.get_pixmap(dpi=150)
        png_bytes = pix.tobytes("png")

        preview_key = f"{preview_prefix}/{page_no}.png"
        minio_client.upload_file(preview_key, png_bytes, content_type="image/png")

        annot_data = json.dumps(
            {"annotations": page_annotations[page_no]},
            ensure_ascii=False,
        )
        annot_key = f"{annot_prefix}/{page_no}.json"
        minio_client.upload_file(
            annot_key, annot_data.encode("utf-8"),
            content_type="application/json",
        )

    pdf_doc.close()
    logger.info(
        f"Documento {doc_id}: {total_pages} previews anotados, "
        f"{annotations_found} correcciones marcadas"
    )
    return total_pages


def _generate_original_page_annotations(
    doc_id: str,
    original_pdf_bytes: bytes,
    all_patches: list[dict],
) -> None:
    """
    Genera anotaciones JSON para las páginas del PDF ORIGINAL, marcando las
    palabras eliminadas o cambiadas respecto al texto corregido.

    Solo produce JSON (sin modificar el PNG original) en annotations_original/.
    """
    pdf_doc = fitz.open(stream=original_pdf_bytes, filetype="pdf")
    total_pages = len(pdf_doc)
    page_annotations: dict[int, list] = {p + 1: [] for p in range(total_pages)}
    total_paragraphs = len(all_patches) or 1

    for patch_idx, patch in enumerate(all_patches):
        orig = patch["original_text"].strip()
        corr = patch["corrected_text"].strip()
        if orig == corr or len(orig) < 3:
            continue

        meta = _get_patch_metadata(patch)
        p_ids = patch.get("patch_ids") or ([patch["patch_id"]] if patch.get("patch_id") else [])

        para_idx = patch.get("paragraph_index", patch_idx)
        est_page = min(int(para_idx / total_paragraphs * total_pages), total_pages - 1)
        search_window = 2
        nearby = list(range(max(0, est_page - search_window),
                            min(total_pages, est_page + search_window + 1)))
        remaining = [p for p in range(total_pages) if p not in set(nearby)]

        # Localizar el párrafo en el PDF original
        found_page_idx: int | None = None
        found_quads: list | None = None

        for max_len in [150, 70, 35]:
            search_text = orig[:max_len] if len(orig) > max_len else orig
            if len(search_text) < 4:
                break
            for page_idx in (nearby + remaining):
                quads = pdf_doc[page_idx].search_for(search_text, quads=True)
                if quads:
                    found_page_idx = page_idx
                    found_quads = list(quads)
                    break
            if found_page_idx is not None:
                break

        if found_page_idx is None or not found_quads:
            continue

        page_no = found_page_idx + 1
        page = pdf_doc[found_page_idx]
        page_rect = page.rect

        ann_base = {
            "patch_ids": p_ids,
            "category": meta["category"],
            "severity": meta["severity"],
            "explanation": meta["explanation"],
            "confidence": patch.get("confidence"),
            "source": patch.get("source", ""),
            "review_status": patch.get("review_status", ""),
            "original_snippet": orig[:100],
            "corrected_snippet": corr[:100],
        }

        # Contorno del párrafo original (sin relleno en el frontend)
        for quad in found_quads:
            r = quad.rect
            page_annotations[page_no].append({
                **ann_base,
                "annot_type": "paragraph",
                "x_pct": round(r.x0 / page_rect.width * 100, 2),
                "y_pct": round(r.y0 / page_rect.height * 100, 2),
                "w_pct": round((r.x1 - r.x0) / page_rect.width * 100, 2),
                "h_pct": round((r.y1 - r.y0) / page_rect.height * 100, 2),
            })

        # Palabras eliminadas/cambiadas (marcadas en rojo en el frontend)
        for phrase in _compute_deleted_phrases_in_original(orig, corr):
            if len(phrase) < 4:
                continue
            del_quads = page.search_for(phrase, quads=True)
            if not del_quads:
                continue
            for dquad in del_quads:
                r = dquad.rect
                page_annotations[page_no].append({
                    **ann_base,
                    "annot_type": "deleted",
                    "x_pct": round(r.x0 / page_rect.width * 100, 2),
                    "y_pct": round(r.y0 / page_rect.height * 100, 2),
                    "w_pct": round((r.x1 - r.x0) / page_rect.width * 100, 2),
                    "h_pct": round((r.y1 - r.y0) / page_rect.height * 100, 2),
                })

    pdf_doc.close()

    for page_no in range(1, total_pages + 1):
        annot_data = json.dumps(
            {"annotations": page_annotations[page_no]},
            ensure_ascii=False,
        )
        annot_key = f"pages/{doc_id}/annotations_original/{page_no}.json"
        minio_client.upload_file(
            annot_key, annot_data.encode("utf-8"),
            content_type="application/json",
        )

    logger.info(f"Documento {doc_id}: anotaciones originales generadas ({total_pages} páginas)")


_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# Fase 6: los helpers de colapso de runs (_clear_run_text_preserve_breaks,
# _copy_run_format), el split proporcional de saltos de página
# (_get_page_break_info, _apply_text_with_page_break, _set_run_text_*_br)
# y la detección parcial de hipervínculos (_get_hyperlink_text_ranges,
# _modification_overlaps_hyperlink) fueron ELIMINADOS: el motor run_splicer
# los reemplaza con edición segura por diff y verificación post-aplicación.


def _apply_text_to_paragraph_runs(paragraph, new_text: str) -> bool:
    """
    Aplica un nuevo texto a un párrafo preservando estructura de formato.

    Fase 6: usa el motor de splice run-level (run_splicer), que edita SOLO
    los w:t afectados por el diff. Preserva negritas/cursivas intra-párrafo,
    hipervínculos, referencias de nota, drawings y saltos (tab/br) por
    construcción, con verificación post-aplicación y rollback.

    El colapso de runs anterior (todo el texto al run dominante) destruía el
    formato interno del párrafo y duplicaba texto de hipervínculos: queda
    eliminado. Si el splice no es aplicable, el único fallback permitido es
    el caso trivial (un solo run de texto puro); en el resto el párrafo se
    omite SIN modificarlo (el patch queda applied=False → revisión).

    Retorna True si hubo cambios aplicados al párrafo.
    """
    old_text = paragraph.text
    if old_text == new_text:
        return False

    from app.config import settings as _settings
    if getattr(_settings, "run_splicer_enabled", True):
        from app.services.run_splicer import splice_paragraph_text
        if splice_paragraph_text(paragraph, new_text):
            return True

    # Fallback seguro: SOLO un run, de texto puro, sin hyperlinks ni
    # elementos especiales. Cualquier otra estructura se omite.
    runs = paragraph.runs
    if (
        len(runs) == 1
        and not paragraph._p.findall('.//' + qn('w:hyperlink'))
        and not runs[0]._r.findall(qn('w:br'))
        and not runs[0]._r.findall(qn('w:tab'))
        and not runs[0]._r.findall(qn('w:drawing'))
        and not runs[0]._r.findall(qn('w:footnoteReference'))
        and (runs[0].text or "") == old_text
    ):
        runs[0].text = new_text
        return True

    logger.info(
        "Párrafo omitido: estructura compleja no editable de forma segura "
        "(sin fallback destructivo)"
    )
    return False


def _get_paragraph_by_location(doc: DocxDocument, location: str):
    """
    Obtiene un párrafo del documento por su ubicación codificada.
    Formatos: 'body:N', 'table:T:R:C:P', 'header:S:P', 'footer:S:P'
    """
    parts = location.split(":")

    if parts[0] == "body":
        idx = int(parts[1])
        if idx < len(doc.paragraphs):
            return doc.paragraphs[idx]

    elif parts[0] == "table":
        t_idx, r_idx, c_idx, p_idx = int(parts[1]), int(parts[2]), int(parts[3]), int(parts[4])
        if t_idx < len(doc.tables):
            table = doc.tables[t_idx]
            if r_idx < len(table.rows):
                row = table.rows[r_idx]
                if c_idx < len(row.cells):
                    cell = row.cells[c_idx]
                    if p_idx < len(cell.paragraphs):
                        return cell.paragraphs[p_idx]

    elif parts[0] in ("header", "footer"):
        s_idx, p_idx = int(parts[1]), int(parts[2])
        if s_idx < len(doc.sections):
            section = doc.sections[s_idx]
            hf = section.header if parts[0] == "header" else section.footer
            if hf and p_idx < len(hf.paragraphs):
                return hf.paragraphs[p_idx]

    return None


# =====================================================================
# Nivel 2 — Sanitización de prefijos de viñeta en patches grupales
# =====================================================================
import re as _re_rendering_group

# Fase 0: SOLO viñetas y numeración decimal. El patrón anterior incluía
# `[a-zA-Z][.)]` y romanos, lo que corrompía texto legítimo:
# "E. coli es…" → "coli es…", "I. Kant sostiene…" → "Kant sostiene…".
_RENDER_LIST_PREFIX = _re_rendering_group.compile(
    r"^\s*(?:[•·▪‒–—●\-\*]|\d{1,3}[.)])\s+"
)


def _strip_list_prefix(text: str, original: str | None = None) -> str:
    """Quita prefijos de viñeta/numeración que el LLM grupal pudo haber dejado.

    Defensivo: el prompt grupal pide "escribe SOLO el texto del ítem", pero
    si el LLM falla no duplicaremos viñetas en el DOCX final.

    Fase 0: si el texto ORIGINAL ya empezaba con el mismo patrón, el prefijo
    es contenido del autor (no un artefacto del LLM) y se preserva.
    """
    if not text:
        return text
    if original and _RENDER_LIST_PREFIX.match(original):
        return text
    return _RENDER_LIST_PREFIX.sub("", text, count=1)


def _apply_individual_patch(doc, patch: dict) -> tuple[bool, str, str]:
    """Aplica un patch individual al DOCX abierto.

    Returns (applied, reason, detail) — reason en {'ok', 'no_paragraph',
    'mismatch'}; detail trae el texto actual del párrafo en mismatch (para log).
    """
    location = patch["location"]
    original_text = patch["original_text"]
    corrected_text = patch["corrected_text"]

    paragraph = _get_paragraph_by_location(doc, location)
    if paragraph is None:
        return False, "no_paragraph", ""

    current_text = paragraph.text.strip()
    if current_text != original_text:
        return False, "mismatch", current_text[:80]

    if _apply_text_to_paragraph_runs(paragraph, corrected_text):
        return True, "ok", ""
    return False, "ok", ""  # no rompió, simplemente sin runs


def _apply_group_patches(doc, group_id, patches: list[dict]) -> tuple[int, int, list[str]]:
    """Aplica un set de patches que pertenecen al mismo grupo.

    Aplica en orden por group_call_index. Si el structural_role indica
    lista MANUAL ("list_item:*:manual"), preserva el prefijo en el texto
    (porque la viñeta es parte del contenido, no del DOCX). En el resto,
    sanitiza prefijos SOLO si el original no los traía (Fase 0).

    Returns (applied, skipped, applied_patch_ids).
    """
    sorted_patches = sorted(
        patches, key=lambda p: p.get("group_call_index", 0) or 0
    )
    applied = 0
    skipped = 0
    applied_ids: list[str] = []
    for patch in sorted_patches:
        patch_clean = dict(patch)
        role = patch.get("structural_role") or ""
        is_manual = role.endswith(":manual")
        if not is_manual:
            patch_clean["corrected_text"] = _strip_list_prefix(
                patch["corrected_text"], original=patch.get("original_text")
            )
        ok, reason, _detail = _apply_individual_patch(doc, patch_clean)
        if ok:
            applied += 1
            applied_ids.extend(patch.get("patch_ids") or [])
            logger.debug(
                f"Grupo {group_id} idx={patch.get('group_call_index')} aplicado"
            )
        else:
            skipped += 1
            logger.warning(
                f"Grupo {group_id} idx={patch.get('group_call_index')}: {reason}"
            )
    return applied, skipped, applied_ids


def _apply_docx_patches(docx_path: str, patches: list[dict]) -> tuple[str, list[str]]:
    """
    Aplica correcciones por párrafo al DOCX original.
    Cada patch tiene {paragraph_index, location, original_text, corrected_text}.

    Para patches con `group_id` no nulo (Nivel 2: corrección grupal de listas
    y tablas), se aplican agrupados y en orden por `group_call_index`, con
    sanitización defensiva de prefijos de viñeta/numeración.

    Verifica que el texto original coincida antes de aplicar.
    Retorna (ruta del DOCX corregido, lista de patch_ids realmente aplicados).
    """
    from collections import defaultdict as _dd
    doc = DocxDocument(docx_path)
    changes_count = 0
    skipped_count = 0
    applied_patch_ids: list[str] = []

    grouped: dict = _dd(list)
    individual: list[dict] = []
    for p in patches:
        if p.get("group_id"):
            grouped[p["group_id"]].append(p)
        else:
            individual.append(p)

    # 1) Grupos primero, para resolver conflictos antes que los individuales
    for gid, gpatches in grouped.items():
        a, s, gids = _apply_group_patches(doc, gid, gpatches)
        changes_count += a
        skipped_count += s
        applied_patch_ids.extend(gids)

    # 2) Patches individuales (path histórico)
    for patch in individual:
        ok, reason, detail = _apply_individual_patch(doc, patch)
        if ok:
            changes_count += 1
            applied_patch_ids.extend(patch.get("patch_ids") or [])
            logger.debug(f"Párrafo {patch['location']} corregido: {patch['source']}")
        else:
            skipped_count += 1
            if reason == "mismatch":
                logger.warning(
                    f"Texto no coincide en {patch['location']}: "
                    f"esperado='{patch['original_text'][:50]}...' "
                    f"actual='{detail}...'"
                )
            else:
                logger.warning(f"No se encontró párrafo en ubicación {patch['location']}")

    # Guardar DOCX corregido
    output_path = str(Path(docx_path).parent / f"{Path(docx_path).stem}_corrected.docx")
    doc.save(output_path)

    logger.info(
        f"DOCX corregido: {changes_count} párrafos modificados, "
        f"{skipped_count} omitidos → {output_path}"
    )
    return output_path, applied_patch_ids, changes_count


def render_docx_first_sync(
    doc_id: str,
    docx_uri: str,
    filename: str,
    all_patches: list[dict],
    docx_bytes_cached: bytes | None = None,
    apply_mode: str = "all",
    render_mode: str = "final",
) -> dict:
    """
    Renderizado Ruta 1: DOCX-first.
    1. Descarga el DOCX original
    2. Aplica correcciones párrafo por párrafo (LanguageTool + GPT)
    3. Genera DOCX corregido
    4. Convierte a PDF con LibreOffice
    5. Genera previews anotados con highlights (MVP2)
    6. Sube ambos a MinIO

    Retorna dict con URIs de los archivos generados.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        # Descargar DOCX original (con cache si disponible)
        local_docx = str(Path(tmpdir) / filename)
        if docx_bytes_cached is not None:
            with open(local_docx, "wb") as _f:
                _f.write(docx_bytes_cached)
        else:
            minio_client.download_file_to_path(docx_uri, local_docx)

        if not all_patches:
            logger.info(f"Documento {doc_id}: sin correcciones que aplicar")
            return {"corrected_docx_uri": None, "corrected_pdf_uri": None, "changes_count": 0}

        # Filtrar patches según apply_mode (Human-in-the-Loop)
        if apply_mode == "accepted_only":
            all_patches = [p for p in all_patches if p.get("review_status") == "accepted"]
        elif apply_mode == "accepted_and_auto":
            all_patches = [
                p for p in all_patches
                if p.get("review_status") in ("accepted", "auto_accepted")
            ]
        # apply_mode == "all" → sin filtro (backward compatible, pipeline original)

        if not all_patches:
            logger.info(f"Documento {doc_id}: sin correcciones aprobadas que aplicar")
            return {"corrected_docx_uri": None, "corrected_pdf_uri": None, "changes_count": 0}

        logger.info(f"Documento {doc_id}: {len(all_patches)} párrafos a corregir (mode={apply_mode})")

        # Aplicar correcciones por párrafo
        corrected_docx_path, applied_patch_ids, applied_count = _apply_docx_patches(
            local_docx, all_patches
        )

        # Convertir DOCX corregido a PDF
        corrected_pdf_path = convert_docx_to_pdf(corrected_docx_path, tmpdir)

        # Generar previews anotados (candidato/final): contorno + palabras cambiadas
        corrected_pdf_bytes = Path(corrected_pdf_path).read_bytes()
        _generate_annotated_previews(doc_id, corrected_pdf_bytes, all_patches, render_mode=render_mode)

        # Generar anotaciones del PDF original: contorno + palabras eliminadas/cambiadas
        stem_inner = Path(filename).stem
        original_pdf_key = f"pdf/{doc_id}/{stem_inner}.pdf"
        try:
            original_pdf_bytes_data = minio_client.download_file(original_pdf_key)
            _generate_original_page_annotations(doc_id, original_pdf_bytes_data, all_patches)
        except Exception as _orig_err:
            logger.warning(f"No se pudieron generar anotaciones originales: {_orig_err}")

        # Subir DOCX corregido a MinIO
        stem = Path(filename).stem
        corrected_docx_key = f"docx/{doc_id}/{stem}_corrected.docx"
        with open(corrected_docx_path, "rb") as f:
            minio_client.upload_file(
                corrected_docx_key, f.read(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # Subir PDF corregido a MinIO
        corrected_pdf_key = f"final/{doc_id}/{stem}_corrected.pdf"
        with open(corrected_pdf_path, "rb") as f:
            minio_client.upload_file(
                corrected_pdf_key, f.read(),
                content_type="application/pdf"
            )

        logger.info(
            f"Documento {doc_id} renderizado: "
            f"DOCX → {corrected_docx_key}, PDF → {corrected_pdf_key}"
        )

        return {
            "corrected_docx_uri": corrected_docx_key,
            "corrected_pdf_uri": corrected_pdf_key,
            # H4: conteo REAL de párrafos aplicados (antes reportaba el total
            # de patches aunque la mitad se hubiera omitido por mismatch).
            "changes_count": applied_count,
            "applied_patch_ids": applied_patch_ids,
        }
