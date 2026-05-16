"""
Sub-etapa B.5 — `extraction_docx`.

Lee el DOCX original con `python-docx` y enriquece los `Block` existentes
(creados por la Etapa B sobre el PDF) con metadatos estructurales:
  - style_name / style_level (Heading 1, etc.)
  - list_id / list_position / list_total / list_format_type / list_level
  - table_id / row_index / column_index / row_total / col_total / table_cell_role
  - element_group_id → FK a `element_groups`

Genera un `ElementGroup` por cada lista (secuencia consecutiva con mismo
numId) y por cada tabla del DOCX. La metadata_json contiene patrones
detectados que el LLM consume en modo grupal (capitalización mayoritaria,
puntuación final mayoritaria, paralelismo estructural).

El matching con `Block` se hace por **location** porque Etapa B ya
produce strings con el mismo esquema que usa `correction.py`
(`body:N`, `table:T:R:C:P`).
"""

from __future__ import annotations

import io
import logging
import re
import uuid
from collections import Counter
from typing import Any

from docx import Document as DocxDocument

from app.models.block import Block
from app.models.element_group import ElementGroup
from app.utils import minio_client

logger = logging.getLogger(__name__)

_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Heurísticas para list_format_type cuando numbering_part no resuelve numFmt
# Se distingue el separador (".", ")") para que listas con marcadores
# inconsistentes (1./2)/3.) caigan en 'mixed'.
_RX_DECIMAL_DOT = re.compile(r"^\s*\d+\.\s+")
_RX_DECIMAL_PAREN = re.compile(r"^\s*\d+\)\s+")
_RX_LOWER_LETTER_DOT = re.compile(r"^\s*[a-z]\.\s+")
_RX_LOWER_LETTER_PAREN = re.compile(r"^\s*[a-z]\)\s+")
_RX_UPPER_LETTER_DOT = re.compile(r"^\s*[A-Z]\.\s+")
_RX_UPPER_LETTER_PAREN = re.compile(r"^\s*[A-Z]\)\s+")
_RX_LOWER_ROMAN = re.compile(r"^\s*[ivxlcdm]+[.)]\s+", re.IGNORECASE)
_RX_BULLET = re.compile(r"^\s*[•·▪‒–—●\-\*]\s+")


# =====================================================================
# Recolección de párrafos del DOCX con metadata nativa
# =====================================================================

def _para_numpr(para) -> tuple[int | None, int | None]:
    """Devuelve (numId, ilvl) leídos del XML w:pPr/w:numPr o (None, None)."""
    try:
        ppr = para._element.pPr
        if ppr is None:
            return None, None
        numPr = ppr.numPr
        if numPr is None:
            return None, None
        numId_el = numPr.numId
        ilvl_el = numPr.ilvl
        num_id = int(numId_el.val) if numId_el is not None else None
        ilvl = int(ilvl_el.val) if ilvl_el is not None else 0
        return num_id, ilvl
    except Exception:
        return None, None


def _para_style_name(para) -> str | None:
    try:
        return para.style.name if para.style is not None else None
    except Exception:
        return None


def _style_level_from_name(name: str | None) -> int | None:
    """Heading 1 → 1, Heading 2 → 2, etc. Cualquier otro → None."""
    if not name:
        return None
    m = re.search(r"(\d)\b", name)
    if m and "head" in name.lower():
        try:
            return int(m.group(1))
        except ValueError:
            return None
    return None


def _collect_paragraphs_with_meta(doc: DocxDocument) -> list[dict]:
    """Recolecta todos los párrafos del DOCX con metadata nativa.

    Devuelve una lista de dicts con:
      text, location, style_name, num_id, ilvl,
      table_idx, row, col, p_idx_in_cell
    """
    items: list[dict] = []

    # Cuerpo
    for i, para in enumerate(doc.paragraphs):
        num_id, ilvl = _para_numpr(para)
        items.append({
            "text": para.text or "",
            "location": f"body:{i}",
            "style_name": _para_style_name(para),
            "num_id": num_id,
            "ilvl": ilvl,
            "table_idx": None,
            "row": None,
            "col": None,
            "p_idx_in_cell": None,
        })

    # Tablas
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    num_id, ilvl = _para_numpr(para)
                    items.append({
                        "text": para.text or "",
                        "location": f"table:{t_idx}:{r_idx}:{c_idx}:{p_idx}",
                        "style_name": _para_style_name(para),
                        "num_id": num_id,
                        "ilvl": ilvl,
                        "table_idx": t_idx,
                        "row": r_idx,
                        "col": c_idx,
                        "p_idx_in_cell": p_idx,
                    })

    # Headers / Footers
    for s_idx, section in enumerate(doc.sections):
        for hf_type, hf in (("header", section.header), ("footer", section.footer)):
            if hf is None:
                continue
            for p_idx, para in enumerate(hf.paragraphs):
                items.append({
                    "text": para.text or "",
                    "location": f"{hf_type}:{s_idx}:{p_idx}",
                    "style_name": _para_style_name(para),
                    "num_id": None,
                    "ilvl": None,
                    "table_idx": None,
                    "row": None,
                    "col": None,
                    "p_idx_in_cell": None,
                })

    return items


# =====================================================================
# Detección de tipo de formato de lista
# =====================================================================

def _resolve_numfmt_from_numbering(doc: DocxDocument, num_id: int) -> str | None:
    """Intenta resolver numFmt desde doc.part.numbering_part.element.

    Devuelve 'decimal'|'lowerLetter'|'upperLetter'|'lowerRoman'|'upperRoman'|'bullet' o None.
    """
    try:
        num_part = doc.part.numbering_part
        if num_part is None:
            return None
        numbering_el = num_part.element
        # Buscar el <w:num w:numId="X"> y obtener abstractNumId
        for num in numbering_el.findall(f"{{{_WORD_NS}}}num"):
            if num.get(f"{{{_WORD_NS}}}numId") != str(num_id):
                continue
            abs_id_el = num.find(f"{{{_WORD_NS}}}abstractNumId")
            if abs_id_el is None:
                return None
            abs_id = abs_id_el.get(f"{{{_WORD_NS}}}val")
            # Buscar <w:abstractNum w:abstractNumId="abs_id">/<w:lvl w:ilvl="0">/<w:numFmt w:val="..."/>
            for abs_num in numbering_el.findall(f"{{{_WORD_NS}}}abstractNum"):
                if abs_num.get(f"{{{_WORD_NS}}}abstractNumId") != abs_id:
                    continue
                # Tomar el primer lvl (ilvl=0) — suficiente para clasificar
                for lvl in abs_num.findall(f"{{{_WORD_NS}}}lvl"):
                    if lvl.get(f"{{{_WORD_NS}}}ilvl") in (None, "0"):
                        fmt_el = lvl.find(f"{{{_WORD_NS}}}numFmt")
                        if fmt_el is not None:
                            return fmt_el.get(f"{{{_WORD_NS}}}val")
        return None
    except Exception as e:
        logger.debug("No se pudo resolver numFmt para numId=%s: %s", num_id, e)
        return None


def _detect_format_from_text(text: str) -> str | None:
    """Distingue separador (".", ")") para detectar listas mixtas.
    Devuelve uno de: decimal_dot, decimal_paren, lowerLetter_dot,
    lowerLetter_paren, upperLetter_dot, upperLetter_paren, lowerRoman, bullet,
    o None.
    """
    if _RX_DECIMAL_DOT.match(text):
        return "decimal_dot"
    if _RX_DECIMAL_PAREN.match(text):
        return "decimal_paren"
    if _RX_UPPER_LETTER_DOT.match(text):
        return "upperLetter_dot"
    if _RX_UPPER_LETTER_PAREN.match(text):
        return "upperLetter_paren"
    if _RX_LOWER_LETTER_DOT.match(text):
        return "lowerLetter_dot"
    if _RX_LOWER_LETTER_PAREN.match(text):
        return "lowerLetter_paren"
    if _RX_LOWER_ROMAN.match(text):
        return "lowerRoman"
    if _RX_BULLET.match(text):
        return "bullet"
    return None


# =====================================================================
# Detección de grupos de lista
# =====================================================================

def _is_list_like_text(text: str) -> bool:
    """True si el texto empieza con un prefijo de lista común
    Y NO parece un título numerado.

    Heurísticas para evitar falsos positivos:
    - "1. Introducción" o "2.1 Objetivos" → títulos, no listas (texto corto,
      sin verbo, sin signos de oración).
    - "1. Esto es una oración larga con verbo." → ítem de lista válido.
    """
    if not text:
        return False
    fmt = _detect_format_from_text(text or "")
    if fmt is None:
        return False
    # Quitar el prefijo para evaluar el resto del texto
    body = text.lstrip()
    # Remover el prefijo "1." / "a)" etc.
    for rx in (
        _RX_DECIMAL_DOT, _RX_DECIMAL_PAREN,
        _RX_LOWER_LETTER_DOT, _RX_LOWER_LETTER_PAREN,
        _RX_UPPER_LETTER_DOT, _RX_UPPER_LETTER_PAREN,
        _RX_LOWER_ROMAN, _RX_BULLET,
    ):
        m = rx.match(body)
        if m:
            body = body[m.end():]
            break
    body = body.strip()
    if len(body) < 4:
        return False
    return True


def _looks_like_numbered_heading(text: str) -> bool:
    """Detecta títulos numerados como '1. Introducción', '2.1 Objetivos',
    'Capítulo 3'. Diferentes de ítems de lista por:
      - Texto corto (< 80 chars)
      - Sin verbo conjugado (heurística simple)
      - Termina sin signo de puntuación final
    """
    s = (text or "").strip()
    if not s:
        return False
    if len(s) > 80:
        return False
    # Termina con puntuación de cierre = no es título
    if s.endswith(".") or s.endswith("?") or s.endswith("!"):
        return False
    return True


def _detect_list_groups(items: list[dict], doc: DocxDocument) -> list[dict]:
    """Detecta listas en el documento por dos caminos:

    1. **Nativa (DOCX numPr):** secuencias consecutivas de párrafos del *body*
       con el mismo `num_id`. Es la señal más fiable.
    2. **Manual (texto):** secuencias consecutivas de párrafos del *body* cuyo
       texto empieza con un prefijo `\\d+[.)]`, `[a-z][.)]`, `[ivx]+[.)]` o
       una viñeta `-`/`•`/`*`/`–`. Captura usuarios que escriben "1." a mano
       sin usar el marcador de lista de Word.

    Devuelve una lista de dicts con:
      num_id, ilvl, item_indexes (índices en `items`),
      format_type, capitalization_majority, ending_punct_majority,
      parallel_structure_hint, detection ('native'|'manual')
    """
    groups: list[dict] = []

    # ── 1) Camino nativo: agrupar por num_id ──
    current: dict | None = None
    used_idxs: set[int] = set()
    for idx, it in enumerate(items):
        nid = it.get("num_id")
        if nid is None:
            if current is not None:
                groups.append(current)
                current = None
            continue
        if current is None or current["num_id"] != nid or current["ilvl"] != it.get("ilvl"):
            if current is not None:
                groups.append(current)
            current = {
                "num_id": nid,
                "ilvl": it.get("ilvl") or 0,
                "item_indexes": [idx],
                "detection": "native",
            }
        else:
            current["item_indexes"].append(idx)
    if current is not None:
        groups.append(current)

    for g in groups:
        used_idxs.update(g["item_indexes"])

    # ── 2) Camino manual: detectar secuencias consecutivas de body con
    #     prefijos de lista. Reglas:
    #       - Mínimo 2 ítems contiguos.
    #       - Cada ítem debe tener cuerpo de texto > 4 chars tras el prefijo.
    #       - Si UN ítem tiene estilo Heading explícito → no es lista (es título).
    #       - Si todos parecen títulos numerados cortos → descartar.
    manual_groups: list[dict] = []
    seq: list[int] = []
    for idx, it in enumerate(items):
        loc = (it.get("location") or "")
        is_body = loc.startswith("body:")
        text = (it.get("text") or "").strip()
        style = (it.get("style_name") or "").lower()
        is_heading_style = "heading" in style or "título" in style or "titulo" in style
        already_used = idx in used_idxs
        eligible = (
            is_body
            and bool(text)
            and not already_used
            and not is_heading_style
            and _is_list_like_text(text)
        )
        if eligible:
            seq.append(idx)
        else:
            if len(seq) >= 2:
                seq_texts = [items[i]["text"] for i in seq]
                # Si todos parecen títulos numerados cortos, no es lista
                all_headings = all(_looks_like_numbered_heading(t) for t in seq_texts)
                if not all_headings:
                    manual_groups.append({
                        "num_id": -len(manual_groups) - 1,
                        "ilvl": 0,
                        "item_indexes": list(seq),
                        "detection": "manual",
                    })
            seq = []
    if len(seq) >= 2:
        seq_texts = [items[i]["text"] for i in seq]
        all_headings = all(_looks_like_numbered_heading(t) for t in seq_texts)
        if not all_headings:
            manual_groups.append({
                "num_id": -len(manual_groups) - 1,
                "ilvl": 0,
                "item_indexes": list(seq),
                "detection": "manual",
            })

    groups.extend(manual_groups)

    # Enriquecer cada grupo con metadata patrón
    for g in groups:
        texts = [items[i]["text"] for i in g["item_indexes"]]
        fmt = None
        if g.get("detection") == "native":
            fmt = _resolve_numfmt_from_numbering(doc, g["num_id"])
        if not fmt and texts:
            detected = [_detect_format_from_text(t) for t in texts if t.strip()]
            distinct = [d for d in detected if d]
            if len(set(distinct)) > 1:
                fmt = "mixed"
            elif distinct:
                fmt = distinct[0] or "bullet"
            else:
                fmt = "bullet"
        g["format_type"] = fmt or "bullet"
        g["capitalization_majority"] = _capitalization_majority(texts)
        g["ending_punct_majority"] = _ending_majority(texts)
        g["parallel_structure_hint"] = _parallel_hint(texts)

    return groups


# =====================================================================
# Patrones para el LLM
# =====================================================================

def _capitalization_majority(texts: list[str]) -> str:
    """'uppercase' si ≥50% empiezan con mayúscula, 'lowercase' si no, 'mixed' si empatado."""
    if not texts:
        return "mixed"
    upper = sum(1 for t in texts if t and t.lstrip()[:1].isupper())
    total = sum(1 for t in texts if t and t.strip())
    if total == 0:
        return "mixed"
    ratio = upper / total
    if ratio >= 0.6:
        return "uppercase"
    if ratio <= 0.4:
        return "lowercase"
    return "mixed"


def _classify_ending(s: str) -> str:
    s = (s or "").rstrip()
    if not s:
        return "empty"
    last = s[-1]
    if last == ".":
        return "period"
    if last in "?¿":
        return "question"
    if last in "!¡":
        return "exclamation"
    if last in ":;,":
        return "punct"
    return "none"


def _ending_majority(texts: list[str]) -> str:
    classes = [_classify_ending(t) for t in texts if t and t.strip()]
    if not classes:
        return "none"
    most, _ = Counter(classes).most_common(1)[0]
    return most


_VERB_INF_RE = re.compile(r"^\s*\w+(ar|er|ir)\b", re.IGNORECASE)
_ARTICLES = {"el", "la", "los", "las", "un", "una", "unos", "unas"}


def _classify_first_word(s: str) -> str:
    first = s.strip().split()[0] if s.strip() else ""
    if _VERB_INF_RE.match(first):
        return "verb_inf"
    if first.lower() in _ARTICLES:
        return "article"
    return "other"


def _parallel_hint(texts: list[str]) -> str:
    items = [t for t in texts if t and t.strip()]
    if len(items) < 2:
        return "consistent"
    classes = [_classify_first_word(t) for t in items]
    most_class, most_count = Counter(classes).most_common(1)[0]
    ratio = most_count / len(classes)
    if ratio >= 0.75:
        return f"consistent:{most_class}"
    return "mixed"


# =====================================================================
# Detección de tablas
# =====================================================================

def _is_header_row(row, doc: DocxDocument) -> bool:
    """Heurística: una fila es header si la mayoría de sus celdas tiene
    al menos un run en bold o un estilo con 'Header' / 'Heading'."""
    n_cells = 0
    bold_cells = 0
    for cell in row.cells:
        any_text = False
        cell_is_bold = False
        for p in cell.paragraphs:
            txt = (p.text or "").strip()
            if not txt:
                continue
            any_text = True
            for run in p.runs:
                if run.bold:
                    cell_is_bold = True
                    break
            sty = _para_style_name(p)
            if sty and ("header" in sty.lower() or "heading" in sty.lower()):
                cell_is_bold = True
                break
            if cell_is_bold:
                break
        if any_text:
            n_cells += 1
            if cell_is_bold:
                bold_cells += 1
    if n_cells == 0:
        return False
    return bold_cells >= max(1, n_cells // 2)


_TOTAL_KEYWORDS = ("total", "suma", "subtotal", "totales")


def _is_totals_row(row) -> bool:
    for cell in row.cells:
        txt = " ".join(p.text or "" for p in cell.paragraphs).strip().lower()
        if any(k in txt for k in _TOTAL_KEYWORDS):
            return True
    return False


_NUM_RX = re.compile(r"^[\-\s]*\$?\s*\d{1,3}(?:[.,]\d{3})*(?:[.,]\d+)?\s*%?\s*$")
_DATE_RX = re.compile(r"\b\d{1,4}[/\-]\d{1,2}[/\-]\d{1,4}\b")


def _column_data_type_hint(samples: list[str]) -> str:
    cleaned = [s.strip() for s in samples if s and s.strip()]
    if not cleaned:
        return "text"
    n_num = sum(1 for s in cleaned if _NUM_RX.match(s))
    n_date = sum(1 for s in cleaned if _DATE_RX.search(s))
    if n_num / len(cleaned) >= 0.6:
        return "number"
    if n_date / len(cleaned) >= 0.5:
        return "date"
    return "text"


def _detect_table_groups(doc: DocxDocument) -> list[dict]:
    """Detecta tablas semánticas del DOCX y descarta cajas/layouts.

    Filtra:
      - Tablas 1×1 (cajas decorativas / cuadros con borde).
      - Tablas Nx1 o 1xN con muy poco contenido (layouts simples).
      - Tablas vacías o con menos de 10 chars totales.
      - Tablas donde solo una celda contiene texto (layouts de columna).
    """
    groups: list[dict] = []
    for t_idx, table in enumerate(doc.tables):
        rows = list(table.rows)
        if not rows:
            continue
        num_rows = len(rows)
        num_cols = max(len(r.cells) for r in rows)

        # 1×1 = caja decorativa
        if num_rows == 1 and num_cols == 1:
            continue
        # Conteo de celdas con contenido
        non_empty = 0
        total_chars = 0
        for r in rows:
            for c in r.cells:
                txt = " ".join((p.text or "") for p in c.paragraphs).strip()
                if txt:
                    non_empty += 1
                    total_chars += len(txt)
        if non_empty == 0 or total_chars < 10:
            continue
        # Layouts con muy pocas celdas significativas — probablemente
        # cuadros decorativos, no tablas de datos.
        if non_empty < 2:
            continue
        # Tablas Nx1 muy pequeñas con poco texto = caja decorativa
        if (num_cols == 1 or num_rows == 1) and total_chars < 80:
            continue
        header_idx = 0 if _is_header_row(rows[0], doc) else None
        totals_idx = (num_rows - 1) if _is_totals_row(rows[-1]) else None

        column_headers: list[str] = []
        if header_idx is not None:
            for c in rows[header_idx].cells[:num_cols]:
                column_headers.append(
                    " ".join(p.text or "" for p in c.paragraphs).strip()
                )

        # Muestra de datos por columna para inferir tipo
        data_hints: list[str] = []
        sample_rows = rows[(header_idx + 1) if header_idx is not None else 0:
                           totals_idx if totals_idx is not None else num_rows]
        for c_idx in range(num_cols):
            samples = []
            for row in sample_rows[:20]:
                if c_idx >= len(row.cells):
                    continue
                samples.append(" ".join(p.text or "" for p in row.cells[c_idx].paragraphs))
            data_hints.append(_column_data_type_hint(samples))

        groups.append({
            "table_idx": t_idx,
            "num_rows": num_rows,
            "num_cols": num_cols,
            "header_row_index": header_idx,
            "totals_row_index": totals_idx,
            "column_headers": column_headers,
            "column_data_type_hints": data_hints,
        })
    return groups


# =====================================================================
# Función pública
# =====================================================================

def extract_docx_structure_sync(
    doc_id: str,
    docx_uri: str,
    session,
    docx_bytes_cached: bytes | None = None,
) -> dict:
    """Punto de entrada de la sub-etapa B.5.

    Args:
        doc_id: UUID del documento (como str)
        docx_uri: clave en MinIO del DOCX original
        session: SQLAlchemy session sync (NO async)
        docx_bytes_cached: bytes del DOCX si ya están en memoria, para evitar
            descargar de MinIO una segunda vez.

    Side-effects:
        - Persiste N filas en `element_groups`.
        - Actualiza filas existentes en `blocks` con metadatos estructurales.

    Returns:
        dict resumen {n_lists, n_tables, n_blocks_updated}
    """
    logger.info("B.5: enriqueciendo estructura DOCX del doc %s", doc_id)
    docx_bytes = docx_bytes_cached or minio_client.download_file(docx_uri)
    doc = DocxDocument(io.BytesIO(docx_bytes))

    items = _collect_paragraphs_with_meta(doc)
    list_groups = _detect_list_groups(items, doc)
    table_groups = _detect_table_groups(doc)

    # 1) Persistir ElementGroup por cada grupo
    list_group_uuid: dict[int, uuid.UUID] = {}  # idx en list_groups → UUID
    for li, g in enumerate(list_groups):
        detection = g.get("detection", "native")
        if detection == "manual":
            docx_native_id = f"manual_list_{li}"
        else:
            docx_native_id = f"numId_{g['num_id']}"
        eg = ElementGroup(
            document_id=doc_id,
            group_type="list",
            docx_native_id=docx_native_id,
            item_count=len(g["item_indexes"]),
            metadata_json={
                "format_type": g["format_type"],
                "level": g["ilvl"],
                "detection": detection,
                "parallel_structure_hint": g["parallel_structure_hint"],
                "ending_punct_majority": g["ending_punct_majority"],
                "capitalization_majority": g["capitalization_majority"],
                "first_item_pos_tag": None,
            },
            correction_status="pending",
        )
        session.add(eg)
        session.flush()
        list_group_uuid[li] = eg.id

    table_group_uuid: dict[int, uuid.UUID] = {}
    for ti, t in enumerate(table_groups):
        eg = ElementGroup(
            document_id=doc_id,
            group_type="table",
            docx_native_id=f"table_{t['table_idx']}",
            item_count=t["num_rows"] * t["num_cols"],
            metadata_json={
                "num_rows": t["num_rows"],
                "num_cols": t["num_cols"],
                "header_row": t["header_row_index"],
                "has_totals_row": t["totals_row_index"] is not None,
                "column_data_type_hints": t["column_data_type_hints"],
                "column_headers": t["column_headers"],
            },
            correction_status="pending",
        )
        session.add(eg)
        session.flush()
        table_group_uuid[ti] = eg.id

    # 2) Indexar items por location para enriquecer Block
    item_by_location: dict[str, dict] = {it["location"]: it for it in items}

    # Mapas auxiliares para list metadata
    list_meta_by_item_idx: dict[int, dict] = {}
    for li, g in enumerate(list_groups):
        total = len(g["item_indexes"])
        for pos, idx in enumerate(g["item_indexes"]):
            list_meta_by_item_idx[idx] = {
                "list_id": f"numId_{g['num_id']}",
                "list_position": pos,
                "list_total": total,
                "list_format_type": g["format_type"],
                "list_level": g["ilvl"],
                "element_group_id": list_group_uuid[li],
            }

    table_meta_by_idx: dict[int, dict] = {ti: t for ti, t in enumerate(table_groups)}

    # 3) Actualizar Block con metadata
    n_updated = 0
    blocks = (
        session.query(Block)
        .join(Block.page)
        .filter_by_doc(doc_id) if False else  # placeholder
        session.query(Block)
    )
    # Filtramos por documento mediante join explícito con Page → Document
    from app.models.page import Page
    blocks = (
        session.query(Block)
        .join(Page, Block.page_id == Page.id)
        .filter(Page.doc_id == doc_id)
        .all()
    )

    # Construir índice item por idx en `items` (mismo orden que items)
    # Esto nos sirve para resolver list_meta a partir de location.
    # location → (idx_in_items)
    idx_by_location: dict[str, int] = {it["location"]: i for i, it in enumerate(items)}

    # Helper: aplica metadata estructural a un Block según su location.
    def _apply_meta(block: Block, loc: str) -> None:
        it = item_by_location.get(loc)
        if it is None:
            return
        block.docx_location = loc
        block.style_name = it.get("style_name")
        block.style_level = _style_level_from_name(it.get("style_name"))
        idx_in_items = idx_by_location.get(loc)
        if idx_in_items is not None and idx_in_items in list_meta_by_item_idx:
            lm = list_meta_by_item_idx[idx_in_items]
            block.list_id = lm["list_id"]
            block.list_position = lm["list_position"]
            block.list_total = lm["list_total"]
            block.list_format_type = lm["list_format_type"]
            block.list_level = lm["list_level"]
            block.element_group_id = lm["element_group_id"]
        if it.get("table_idx") is not None:
            ti = it["table_idx"]
            t = table_meta_by_idx.get(ti)
            if t is not None:
                block.table_id = f"table_{ti}"
                block.row_index = it.get("row")
                block.column_index = it.get("col")
                block.row_total = t["num_rows"]
                block.col_total = t["num_cols"]
                if t["header_row_index"] is not None and it.get("row") == t["header_row_index"]:
                    block.table_cell_role = "header"
                elif t["totals_row_index"] is not None and it.get("row") == t["totals_row_index"]:
                    block.table_cell_role = "total"
                else:
                    block.table_cell_role = "data"
                block.element_group_id = table_group_uuid.get(ti)

    # 1ª pasada: matchear cada Block existente con la mejor location DOCX
    matched_locations: set[str] = set()
    for block in blocks:
        loc = _guess_location_for_block(block, items)
        if loc is None or loc in matched_locations:
            continue
        matched_locations.add(loc)
        _apply_meta(block, loc)
        n_updated += 1

    # 2ª pasada: crear Block SINTÉTICOS para items DOCX que pertenecen a un
    # grupo (lista/tabla) y no fueron matcheados. PyMuPDF a veces colapsa
    # celdas de tabla en menos bloques visuales que celdas hay, por lo que
    # muchas celdas no encuentran Block. Los Block sintéticos permiten que
    # la pasada D.5 grupal y el rendering operen sobre cada celda/ítem.
    locations_in_groups: set[str] = set()
    for idx_in_items in list_meta_by_item_idx.keys():
        if 0 <= idx_in_items < len(items):
            locations_in_groups.add(items[idx_in_items]["location"])
    for it in items:
        if it.get("table_idx") is not None:
            locations_in_groups.add(it["location"])

    missing_locations = [
        loc for loc in locations_in_groups
        if loc not in matched_locations
    ]

    if missing_locations:
        # Elegimos una página representativa (la primera) para colgar los
        # blocks sintéticos. El rendering los localiza por docx_location,
        # no por page_id, así que la asignación es solo administrativa.
        first_page = (
            session.query(Page)
            .filter(Page.doc_id == doc_id)
            .order_by(Page.page_no.asc())
            .first()
        )
        if first_page is not None:
            # block_no inicial = max actual + 1 para evitar conflictos
            max_block_no = (
                session.query(Block.block_no)
                .filter(Block.page_id == first_page.id)
                .order_by(Block.block_no.desc())
                .first()
            )
            next_block_no = (max_block_no[0] if max_block_no else 0) + 1

            for loc in missing_locations:
                it = item_by_location.get(loc)
                if it is None:
                    continue
                text = it.get("text") or ""
                if not text.strip():
                    continue
                synthetic = Block(
                    page_id=first_page.id,
                    block_no=next_block_no,
                    block_type="docx_synthetic",
                    bbox_x0=0.0, bbox_y0=0.0, bbox_x1=0.0, bbox_y1=0.0,
                    original_text=text,
                    requires_llm=True,
                )
                _apply_meta(synthetic, loc)
                session.add(synthetic)
                next_block_no += 1
                n_updated += 1
            session.flush()

    session.commit()
    summary = {
        "n_lists": len(list_groups),
        "n_tables": len(table_groups),
        "n_blocks_updated": n_updated,
        "n_synthetic_blocks": len(missing_locations),
    }
    logger.info("B.5 completado: %s", summary)
    return summary


_NORM_WS = re.compile(r"\s+")


def _normalize_for_match(text: str | None) -> str:
    """Colapsa espacios y lowercase para matching tolerante."""
    if not text:
        return ""
    return _NORM_WS.sub(" ", text).strip().lower()


def _guess_location_for_block(block: Block, items: list[dict]) -> str | None:
    """Resuelve la location DOCX para un Block dado.

    Estrategia (en orden):
      1) Match exacto por texto normalizado (whitespace + case).
      2) Match por prefijo de 80 caracteres normalizado (tolera fragmentaciones
         del PDF que truncan o agregan líneas extra).
      3) Match por substring contained (uno contiene al otro).

    Es un best-effort: si no se encuentra, devuelve None y el block queda sin
    enriquecer (retrocompatible).
    """
    target = (block.original_text or "").strip()
    if not target:
        return None
    norm_target = _normalize_for_match(target)
    if not norm_target:
        return None

    norm_items = [(_normalize_for_match(it.get("text")), it["location"]) for it in items]

    # Exact normalized
    for nt, loc in norm_items:
        if nt and nt == norm_target:
            return loc

    # Prefix (80 chars) — útil cuando el PDF fragmenta líneas
    prefix_target = norm_target[:80]
    if len(prefix_target) >= 20:
        for nt, loc in norm_items:
            if nt and (nt.startswith(prefix_target) or prefix_target.startswith(nt[:80])):
                return loc

    # Containment — para fragmentaciones agresivas
    for nt, loc in norm_items:
        if nt and len(nt) >= 30 and (nt in norm_target or norm_target in nt):
            return loc

    return None
