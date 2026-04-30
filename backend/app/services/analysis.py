"""
Servicio de análisis editorial — Etapa C del pipeline (MVP2 Lote 3).

Se ejecuta UNA VEZ por documento antes de cualquier corrección.
Sub-etapas:
  C.1: Inferencia de género/audiencia/registro
  C.2: Completar/validar el style profile
  C.3: Generar resúmenes por sección
  C.4: Extraer glosario y términos protegidos
  C.5: Clasificar párrafos por tipo
"""

import json
import logging
import re
import tempfile
from collections import Counter
from pathlib import Path

from docx import Document as DocxDocument

from app.config import settings
from app.utils import minio_client
from app.utils.openai_client import openai_client, _extract_usage

logger = logging.getLogger(__name__)


# =============================================
# C.5: Clasificación heurística de párrafos
# =============================================

# Patrones para detección de listas
_LIST_PATTERN = re.compile(
    r"^\s*(?:[•\-–—\*]|\d{1,3}[.)]\s|[a-zA-Z][.)]\s|[ivxIVX]+[.)]\s)"
)

# Patrones para detección de diálogo
_DIALOGUE_PATTERN = re.compile(r"^\s*[—–\-]\s|^\"[A-ZÁÉÍÓÚÑ]|^«")


def classify_paragraph(
    text: str,
    location: str,
    style_name: str | None = None,
    is_in_table: bool = False,
    glossary_terms: list[str] | None = None,
) -> tuple[str, bool]:
    """
    Clasifica un párrafo por tipo usando heurísticas (sin LLM).

    Args:
        text: Texto del párrafo.
        location: Ubicación (body:N, table:T:R:C:P, header:S:P, footer:S:P).
        style_name: Nombre del estilo python-docx (ej: 'Heading 1').
        is_in_table: Si el párrafo está dentro de una tabla.
        glossary_terms: Lista de términos técnicos del glosario.

    Returns:
        (paragraph_type, requires_llm)
    """
    stripped = text.strip()
    if not stripped:
        return "vacio", False

    # Por ubicación en documento
    if location.startswith("header:"):
        return "encabezado", False
    if location.startswith("footer:"):
        return "footer", False
    if is_in_table or location.startswith("table:"):
        return "celda_tabla", True

    # Por estilo de python-docx
    if style_name:
        sl = style_name.lower()
        if "heading" in sl or "título" in sl or "titulo" in sl:
            if "1" in sl:
                return "titulo", False
            return "subtitulo", False
        if "quote" in sl or "cita" in sl:
            return "cita", False
        if "list" in sl or "lista" in sl:
            return "lista", True
        if "title" in sl:
            return "titulo", False

    # Por contenido
    if _DIALOGUE_PATTERN.match(stripped):
        return "dialogo", True

    if _LIST_PATTERN.match(stripped):
        return "lista", True

    # Pies de figura/tabla: SOLO patrones específicos de caption (no startswith genérico).
    # Falsos positivos a evitar: "Tabla de contenidos", "Figura retórica", "Tabla comparativa de X",
    # cualquier oración corta legítima de < 20 chars.
    # Match estricto: "Figura 1.", "Fig. 2:", "Tabla 3 -", "Fuente:", "Nota:", "Elaboración propia"
    stripped_lower = stripped.lower()
    is_caption_numbered = bool(
        re.match(
            r'^(?:figura|fig\.?|tabla|cuadro|imagen|gráfico|grafico|mapa|esquema|diagrama)\s+\d+\s*[\.\:\-–—]',
            stripped_lower,
        )
    )
    is_caption_attribution = bool(
        re.match(r'^(?:fuente|nota|elaboraci[oó]n propia)\s*[\:\.\-]', stripped_lower)
    )
    if is_caption_numbered or is_caption_attribution:
        return "pie_imagen", True

    # Detectar densidad de términos técnicos
    if glossary_terms and len(glossary_terms) > 0:
        lower_text = stripped.lower()
        term_count = sum(1 for t in glossary_terms if t.lower() in lower_text)
        words = len(stripped.split())
        if words > 0 and term_count / words > 0.1:
            return "explicacion_tecnica", True

    # Default
    return "narrativo", True


# =============================================
# C.3: Detección de secciones
# =============================================

def _detect_sections(
    doc: DocxDocument,
    all_paragraphs: list[tuple[str, str]],
) -> list[dict]:
    """
    Detecta secciones en el DOCX basándose en estilos Heading.

    Returns:
        Lista de dicts: {section_index, section_title, start_paragraph, end_paragraph}
    """
    sections = []
    current_section = {
        "section_index": 0,
        "section_title": None,
        "start_paragraph": 0,
    }

    body_paragraphs = doc.paragraphs
    body_idx = 0

    for global_idx, (text, location) in enumerate(all_paragraphs):
        if not location.startswith("body:"):
            continue

        # Buscar si el párrafo del cuerpo es un heading
        if body_idx < len(body_paragraphs):
            para = body_paragraphs[body_idx]
            style_name = para.style.name if para.style else ""

            if style_name and ("heading" in style_name.lower() or "título" in style_name.lower()):
                # Cerrar sección anterior
                if global_idx > current_section["start_paragraph"]:
                    current_section["end_paragraph"] = global_idx - 1
                    sections.append(current_section.copy())

                # Abrir nueva sección
                current_section = {
                    "section_index": len(sections),
                    "section_title": text.strip() if text.strip() else None,
                    "start_paragraph": global_idx,
                }

        body_idx += 1

    # Cerrar última sección
    current_section["end_paragraph"] = len(all_paragraphs) - 1
    sections.append(current_section)

    # Si no se detectaron headings, crear secciones cada ~30 párrafos
    if len(sections) <= 1 and len(all_paragraphs) > 40:
        sections = []
        chunk_size = 30
        for i in range(0, len(all_paragraphs), chunk_size):
            end = min(i + chunk_size - 1, len(all_paragraphs) - 1)
            sections.append({
                "section_index": len(sections),
                "section_title": None,
                "start_paragraph": i,
                "end_paragraph": end,
            })

    return sections


# =============================================
# C.4: Extracción de términos
# =============================================

def _extract_terms(
    all_paragraphs: list[tuple[str, str]],
    profile_protected: list[str] | None = None,
) -> list[dict]:
    """
    Extrae términos técnicos recurrentes del documento
    usando análisis de frecuencia de n-gramas.
    """
    # Recopilar todo el texto del cuerpo
    full_text = " ".join(
        text for text, loc in all_paragraphs
        if loc.startswith("body:") and text.strip()
    )

    words = full_text.split()
    if len(words) < 20:
        return []

    # Buscar bigramas y trigramas recurrentes
    term_counter: Counter = Counter()
    first_occurrence: dict[str, int] = {}

    for para_idx, (text, _loc) in enumerate(all_paragraphs):
        text_words = text.strip().split()
        for n in (2, 3):
            for i in range(len(text_words) - n + 1):
                ngram = " ".join(text_words[i:i+n])
                # Filtrar n-gramas que son solo stopwords o muy cortos
                ngram_clean = ngram.strip(".,;:()\"'«»")
                if len(ngram_clean) < 5:
                    continue
                # Ignorar n-gramas con solo minúsculas comunes
                ngram_lower = ngram_clean.lower()
                term_counter[ngram_lower] += 1
                if ngram_lower not in first_occurrence:
                    first_occurrence[ngram_lower] = para_idx

    # Filtrar: frecuencia >= 3, longitud razonable
    MIN_FREQ = 3
    terms = []
    seen_substrings = set()

    for term, freq in term_counter.most_common(100):
        if freq < MIN_FREQ:
            continue
        if len(term) < 6:
            continue
        # Evitar duplicados por substring
        if any(term in s or s in term for s in seen_substrings):
            continue
        seen_substrings.add(term)

        terms.append({
            "term": term,
            "normalized_form": term,
            "frequency": freq,
            "first_occurrence_paragraph": first_occurrence.get(term, 0),
            "is_protected": False,
            "decision": "use_as_is",
        })

    # Agregar términos protegidos del perfil
    if profile_protected:
        existing = {t["term"].lower() for t in terms}
        for pt in profile_protected:
            pt_lower = pt.lower()
            if pt_lower not in existing:
                terms.append({
                    "term": pt,
                    "normalized_form": pt,
                    "frequency": 0,
                    "first_occurrence_paragraph": 0,
                    "is_protected": True,
                    "decision": "use_as_is",
                })
            else:
                # Marcar como protegido si ya existe
                for t in terms:
                    if t["term"].lower() == pt_lower:
                        t["is_protected"] = True
                        break

    return terms[:50]  # Limitar a 50 términos


# =============================================
# C.1 + C.2: Inferencia de perfil con LLM
# =============================================

def _infer_profile_with_llm(
    sample_paragraphs: list[str],
    profile: dict | None = None,
) -> tuple[dict, dict]:
    """
    Analiza muestras del documento con LLM para inferir
    características editoriales faltantes.

    Returns:
        (inferred_dict, usage_dict)
    """
    empty_usage = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}

    if not openai_client.client:
        logger.warning("OpenAI no disponible, usando inferencia por defecto")
        return {
            "genre": "no_ficcion_general",
            "audience_type": "adultos",
            "register": "neutro",
            "tone": "neutro",
            "spanish_variant": "es",
            "key_terms": [],
            "suggested_intervention": "moderada",
        }, empty_usage

    # Construir muestra compacta
    sample_text = "\n---\n".join(p[:300] for p in sample_paragraphs[:8])

    system_msg = (
        "Eres un editor profesional. Analiza el extracto de un documento en español "
        "y clasifícalo. Responde SOLO con JSON válido."
    )

    user_msg = f"""Analiza este extracto y responde en JSON:

EXTRACTO:
{sample_text}

RESPONDE CON ESTE FORMATO JSON:
{{
  "genre": "infantil|juvenil|novela_literaria|ensayo|psicologia_academica|psicologia_divulgativa|manual_tecnico|texto_marketing|no_ficcion_general|otro",
  "audience_type": "niños_6_8|niños_9_12|adolescentes|adultos_no_especialistas|adultos_especialistas|profesionales|general",
  "register": "informal_claro|neutro_claro|neutro|formal_claro|formal_tecnico|persuasivo",
  "tone": "reflexivo|didactico|narrativo|persuasivo|neutro|ludico",
  "spanish_variant": "es-ES|es-MX|es-CO|es-AR|es",
  "key_terms": ["término1", "término2", "...hasta 10 términos técnicos o recurrentes"],
  "suggested_intervention": "minima|sutil|moderada|agresiva"
}}"""

    try:
        response = openai_client.client.chat.completions.create(
            model=settings.openai_model,
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": user_msg},
            ],
            max_completion_tokens=300,
            temperature=0.2,
            response_format={"type": "json_object"},
        )
        content = response.choices[0].message.content
        usage = _extract_usage(response)
        data = json.loads(content)
        logger.info(f"C.1: Perfil inferido — género={data.get('genre')}, registro={data.get('register')}")
        return data, usage

    except Exception as e:
        logger.error(f"Error en inferencia de perfil: {e}")
        return {
            "genre": "no_ficcion_general",
            "audience_type": "adultos",
            "register": "neutro",
            "tone": "neutro",
            "spanish_variant": "es",
            "key_terms": [],
            "suggested_intervention": "moderada",
        }, empty_usage


# =============================================
# C.3: Resúmenes por sección con LLM
# =============================================

def _summarize_sections_with_llm(
    sections: list[dict],
    all_paragraphs: list[tuple[str, str]],
) -> tuple[list[dict], dict]:
    """
    Genera resúmenes para cada sección usando una llamada batch al LLM.

    Returns:
        (sections_with_summaries, usage_dict)
    """
    empty_usage = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}

    if not openai_client.client:
        # Sin LLM: llenar con datos básicos
        for sec in sections:
            sec["summary_text"] = None
            sec["topic"] = sec.get("section_title", f"Sección {sec['section_index'] + 1}")
            sec["local_tone"] = None
            sec["active_terms"] = []
            sec["transition_from_previous"] = None
        return sections, empty_usage

    # Construir descripción compacta de cada sección
    section_descriptions = []
    for sec in sections:
        start = sec["start_paragraph"]
        end = sec["end_paragraph"]
        # Tomar primer y último párrafo de la sección como muestra
        section_texts = []
        for idx in range(start, min(end + 1, len(all_paragraphs))):
            text = all_paragraphs[idx][0].strip()
            if text and len(text) > 10:
                section_texts.append(text[:200])
            if len(section_texts) >= 4:
                break

        title = sec.get("section_title") or f"Sección {sec['section_index'] + 1}"
        sample = " | ".join(section_texts) if section_texts else "(vacía)"
        section_descriptions.append(f"[{title}]: {sample}")

    # Limitar a 15 secciones para no exceder tokens
    section_descriptions = section_descriptions[:15]

    system_msg = (
        "Eres un editor profesional. Analiza las secciones de un documento "
        "y genera resúmenes breves. Responde SOLO con JSON válido."
    )

    user_msg = f"""Para cada sección, genera un resumen breve (~30 palabras), tema principal, tono local, y términos activos.

SECCIONES:
{chr(10).join(f"{i}. {desc}" for i, desc in enumerate(section_descriptions))}

RESPONDE CON ESTE FORMATO JSON:
{{
  "sections": [
    {{
      "summary": "resumen de ~30 palabras",
      "topic": "tema principal",
      "tone": "reflexivo|didactico|narrativo|persuasivo|neutro",
      "terms": ["término1", "término2"],
      "transition": "cómo conecta con la sección anterior o null"
    }}
  ]
}}"""

    try:
        response = openai_client.client.chat.completions.create(
            model=settings.openai_model,
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": user_msg},
            ],
            max_completion_tokens=800,
            temperature=0.2,
            response_format={"type": "json_object"},
        )
        content = response.choices[0].message.content
        usage = _extract_usage(response)
        data = json.loads(content)
        summaries = data.get("sections", [])

        for i, sec in enumerate(sections):
            if i < len(summaries):
                s = summaries[i]
                sec["summary_text"] = s.get("summary")
                sec["topic"] = s.get("topic", sec.get("section_title"))
                sec["local_tone"] = s.get("tone")
                sec["active_terms"] = s.get("terms", [])
                sec["transition_from_previous"] = s.get("transition")
            else:
                sec["summary_text"] = None
                sec["topic"] = sec.get("section_title")
                sec["local_tone"] = None
                sec["active_terms"] = []
                sec["transition_from_previous"] = None

        logger.info(f"C.3: {len(summaries)} secciones resumidas")
        return sections, usage

    except Exception as e:
        logger.error(f"Error generando resúmenes de secciones: {e}")
        for sec in sections:
            sec["summary_text"] = None
            sec["topic"] = sec.get("section_title")
            sec["local_tone"] = None
            sec["active_terms"] = []
            sec["transition_from_previous"] = None
        return sections, empty_usage


# =============================================
# Función principal: analyze_document_sync
# =============================================

def _collect_all_paragraphs_with_styles(doc: DocxDocument) -> list[dict]:
    """
    Recolecta todos los párrafos del DOCX con su ubicación y estilo.
    Similar a correction.py._collect_all_paragraphs pero incluye metadatos.
    """
    paragraphs = []

    # Cuerpo principal
    for i, para in enumerate(doc.paragraphs):
        paragraphs.append({
            "text": para.text,
            "location": f"body:{i}",
            "style_name": para.style.name if para.style else None,
            "is_in_table": False,
        })

    # Tablas
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    paragraphs.append({
                        "text": para.text,
                        "location": f"table:{t_idx}:{r_idx}:{c_idx}:{p_idx}",
                        "style_name": para.style.name if para.style else None,
                        "is_in_table": True,
                    })

    # Headers y footers
    for s_idx, section in enumerate(doc.sections):
        for hf_type, hf in [("header", section.header), ("footer", section.footer)]:
            if hf is None:
                continue
            for p_idx, para in enumerate(hf.paragraphs):
                paragraphs.append({
                    "text": para.text,
                    "location": f"{hf_type}:{s_idx}:{p_idx}",
                    "style_name": para.style.name if para.style else None,
                    "is_in_table": False,
                })

    return paragraphs


def analyze_document_sync(
    doc_id: str,
    docx_uri: str,
    profile: dict | None = None,
    docx_bytes_cached: bytes | None = None,
) -> dict:
    """
    Etapa C: Análisis editorial del documento.
    Se ejecuta UNA VEZ antes de cualquier corrección.

    Args:
        doc_id: UUID del documento.
        docx_uri: Ruta en MinIO del DOCX original.
        profile: Dict con campos del perfil editorial (o None).

    Returns:
        dict con: inferred_profile, sections, terms, paragraph_classifications,
                  usage_records, stats
    """
    logger.info(f"[Etapa C] Iniciando análisis editorial para {doc_id}")

    usage_records = []

    # Descargar DOCX (con cache si disponible)
    if docx_bytes_cached is not None:
        docx_bytes = docx_bytes_cached
    else:
        docx_bytes = minio_client.download_file(docx_uri)
    tmpfile = tempfile.mktemp(suffix=".docx")
    with open(tmpfile, "wb") as f:
        f.write(docx_bytes)

    doc = DocxDocument(tmpfile)
    Path(tmpfile).unlink(missing_ok=True)

    # Recolectar párrafos con metadatos
    all_paras = _collect_all_paragraphs_with_styles(doc)
    all_paragraphs_simple = [(p["text"], p["location"]) for p in all_paras]
    logger.info(f"[Etapa C] Documento tiene {len(all_paras)} párrafos totales")

    # =============================================
    # C.1 + C.2: Inferencia de perfil
    # =============================================
    sample_texts = [
        p["text"] for p in all_paras
        if p["location"].startswith("body:") and len(p["text"].strip()) > 20
    ]
    inferred_profile, infer_usage = _infer_profile_with_llm(
        sample_paragraphs=sample_texts,
        profile=profile,
    )
    if infer_usage["total_tokens"] > 0:
        _cost = (
            infer_usage["prompt_tokens"] / 1e6 * settings.openai_pricing_input
            + infer_usage["completion_tokens"] / 1e6 * settings.openai_pricing_output
        )
        usage_records.append({
            "paragraph_index": -1,
            "location": "analysis:profile_inference",
            "call_type": "analysis_c1",
            "model_used": settings.openai_model,
            **infer_usage,
            "cost_usd": round(_cost, 8),
        })

    # Merge: perfil inferido completa campos faltantes del perfil explícito
    profile_updates = {}
    if profile and inferred_profile:
        for key in ("genre", "audience_type", "register", "tone"):
            if not profile.get(key) and inferred_profile.get(key):
                profile_updates[key] = inferred_profile[key]

    # =============================================
    # C.4: Extracción de términos
    # =============================================
    profile_protected = (profile or {}).get("protected_terms", [])
    llm_terms = inferred_profile.get("key_terms", [])
    # Combinar términos del perfil + inferidos
    combined_protected = list(set(profile_protected + llm_terms))

    terms = _extract_terms(all_paragraphs_simple, combined_protected)
    # Marcar términos inferidos por LLM como protegidos
    for t in terms:
        if t["term"].lower() in {x.lower() for x in llm_terms}:
            t["is_protected"] = True

    glossary_terms = [t["term"] for t in terms]
    logger.info(f"[Etapa C] C.4: {len(terms)} términos extraídos ({sum(1 for t in terms if t['is_protected'])} protegidos)")

    # =============================================
    # C.3: Detección de secciones + resúmenes
    # =============================================
    sections = _detect_sections(doc, all_paragraphs_simple)
    sections, section_usage = _summarize_sections_with_llm(sections, all_paragraphs_simple)
    if section_usage["total_tokens"] > 0:
        _cost = (
            section_usage["prompt_tokens"] / 1e6 * settings.openai_pricing_input
            + section_usage["completion_tokens"] / 1e6 * settings.openai_pricing_output
        )
        usage_records.append({
            "paragraph_index": -1,
            "location": "analysis:section_summaries",
            "call_type": "analysis_c3",
            "model_used": settings.openai_model,
            **section_usage,
            "cost_usd": round(_cost, 8),
        })
    logger.info(f"[Etapa C] C.3: {len(sections)} secciones detectadas")

    # =============================================
    # C.5: Clasificación de párrafos
    # =============================================
    paragraph_classifications = []
    type_counter: Counter = Counter()

    for idx, p in enumerate(all_paras):
        text = p["text"].strip()
        if not text:
            continue

        ptype, needs_llm = classify_paragraph(
            text=text,
            location=p["location"],
            style_name=p["style_name"],
            is_in_table=p["is_in_table"],
            glossary_terms=glossary_terms,
        )

        paragraph_classifications.append({
            "paragraph_index": idx,
            "location": p["location"],
            "paragraph_type": ptype,
            "requires_llm": needs_llm,
            "text_preview": text[:80],
        })
        type_counter[ptype] += 1

    logger.info(f"[Etapa C] C.5: Clasificación — {dict(type_counter)}")

    # Estadísticas
    stats = {
        "total_paragraphs": len(all_paras),
        "non_empty_paragraphs": len(paragraph_classifications),
        "sections_detected": len(sections),
        "terms_extracted": len(terms),
        "terms_protected": sum(1 for t in terms if t["is_protected"]),
        "paragraph_types": dict(type_counter),
        "paragraphs_needing_llm": sum(1 for c in paragraph_classifications if c["requires_llm"]),
        "analysis_llm_calls": len(usage_records),
        "analysis_total_tokens": sum(r["total_tokens"] for r in usage_records),
        "analysis_total_cost": round(sum(r["cost_usd"] for r in usage_records), 6),
    }

    logger.info(
        f"[Etapa C] Análisis completo: {stats['sections_detected']} secciones, "
        f"{stats['terms_extracted']} términos, "
        f"{stats['paragraphs_needing_llm']}/{stats['non_empty_paragraphs']} párrafos necesitan LLM, "
        f"costo: ${stats['analysis_total_cost']:.6f}"
    )

    return {
        "inferred_profile": inferred_profile,
        "profile_updates": profile_updates,
        "sections": sections,
        "terms": terms,
        "paragraph_classifications": paragraph_classifications,
        "usage_records": usage_records,
        "stats": stats,
    }


# =============================================
# C.6: Análisis de Contexto Global (Plan v4)
# =============================================

def analyze_global_context_sync(
    doc_id: str,
    all_paragraphs: list[tuple[str, str]],
    profile: dict | None = None,
    protected_terms: list[str] | None = None,
) -> dict:
    """
    Etapa C.6: genera el "ADN editorial" del documento mediante muestreo estratificado.

    Args:
        doc_id: UUID del documento.
        all_paragraphs: Lista de (text, location) de todos los párrafos.
        profile: Perfil editorial (para inyectar en el prompt).
        protected_terms: Términos ya protegidos detectados en C.4.

    Returns:
        dict con: global_summary, dominant_voice, dominant_register,
                  key_themes_json, protected_globals_json, style_fingerprint_json,
                  usage_record
    """
    logger.info(f"[Etapa C.6] Generando contexto global para {doc_id}")

    # Muestreo estratificado: primeros 3, medianos 3, últimos 3
    body_paras = [
        (text, loc) for text, loc in all_paragraphs
        if loc.startswith("body:") and len(text.strip()) > 30
    ]
    n = len(body_paras)

    sample = []
    if n > 0:
        # Primeros
        sample.extend(body_paras[:3])
        # Medianos
        if n > 6:
            mid = n // 2
            sample.extend(body_paras[max(0, mid - 1):mid + 2])
        # Últimos
        if n > 3:
            sample.extend(body_paras[-3:])

    # Deduplicar manteniendo orden
    seen = set()
    unique_sample = []
    for p in sample:
        if p[0] not in seen:
            seen.add(p[0])
            unique_sample.append(p)
    sample = unique_sample[:9]  # máx 9 párrafos

    if not sample:
        logger.warning(f"[Etapa C.6] No hay párrafos de cuerpo para muestrear en {doc_id}")
        return {
            "global_summary": None,
            "dominant_voice": None,
            "dominant_register": None,
            "key_themes_json": [],
            "protected_globals_json": [],
            "style_fingerprint_json": {},
            "usage_record": None,
        }

    sample_text = "\n\n---\n\n".join(f"[{loc}]\n{text}" for text, loc in sample)

    known_protected = (protected_terms or [])[:20]
    profile_hint = ""
    if profile:
        genre = profile.get("genre", "")
        register = profile.get("register", "")
        if genre or register:
            profile_hint = f"\nPerfil conocido: género={genre}, registro={register}."

    prompt = f"""Analiza los siguientes fragmentos representativos de un documento y genera su "ADN editorial".{profile_hint}
Términos ya detectados como protegidos: {', '.join(known_protected) if known_protected else 'ninguno aún'}.

FRAGMENTOS DEL DOCUMENTO:
{sample_text}

Responde SOLO con este JSON:
{{
  "global_summary": "resumen global del documento en ~200 palabras",
  "dominant_voice": "descripción de la voz y estilo del autor en ~60 palabras",
  "dominant_register": "academico_formal|divulgativo|narrativo_literario|tecnico|periodistico|corporativo|otro",
  "key_themes": [{{"theme": "nombre del tema", "weight": 0.0-1.0}}],
  "protected_globals": [{{"term": "término exacto", "reason": "por qué protegerlo"}}],
  "style_fingerprint": {{
    "avg_sentence_length": número_aproximado_palabras,
    "passive_voice_ratio": 0.0-1.0,
    "uses_dashes": true|false,
    "uses_parentheses": true|false,
    "formality_score": 0.0-1.0
  }}
}}"""

    system = (
        "Eres un analista editorial experto en español. "
        "Analizas la voz, estilo y temática de textos literarios y académicos. "
        "Responde siempre con JSON válido."
    )

    empty_usage = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}

    data, usage = openai_client.correct_with_profile(
        system_prompt=system,
        user_prompt=prompt,
        max_length=None,
        model_override=settings.openai_editorial_model,
        max_tokens_override=settings.openai_audit_max_tokens,
    )

    if data is None:
        logger.warning(f"[Etapa C.6] LLM no respondió para {doc_id}")
        return {
            "global_summary": None,
            "dominant_voice": None,
            "dominant_register": None,
            "key_themes_json": [],
            "protected_globals_json": [],
            "style_fingerprint_json": {},
            "usage_record": None,
        }

    usage_record = None
    if usage.get("total_tokens", 0) > 0:
        _cost = (
            usage["prompt_tokens"] / 1e6 * settings.openai_pricing_input
            + usage["completion_tokens"] / 1e6 * settings.openai_pricing_output
        )
        usage_record = {
            "paragraph_index": -2,
            "location": "analysis:global_context",
            "call_type": "analysis_c6",
            "model_used": settings.openai_editorial_model,
            **usage,
            "cost_usd": round(_cost, 8),
        }

    logger.info(
        f"[Etapa C.6] Contexto global generado: "
        f"register={data.get('dominant_register')}, "
        f"términos_protegidos={len(data.get('protected_globals', []))}, "
        f"tokens={usage.get('total_tokens', 0)}"
    )

    return {
        "global_summary": data.get("global_summary"),
        "dominant_voice": data.get("dominant_voice"),
        "dominant_register": data.get("dominant_register"),
        "key_themes_json": data.get("key_themes", []),
        "protected_globals_json": data.get("protected_globals", []),
        "style_fingerprint_json": data.get("style_fingerprint", {}),
        "usage_record": usage_record,
    }
