"""
Tests unitarios para _collect_all_paragraphs y detección de w:br type="page".
"""

import io
import pytest
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from lxml import etree


def _make_doc_with_page_break() -> DocxDocument:
    """Crea un documento DOCX con un párrafo que contiene salto de página interno."""
    doc = DocxDocument()
    para = doc.add_paragraph("Texto antes del salto")
    # Añadir w:br type="page" al run
    run = para.runs[0]
    br_elem = etree.SubElement(run._r, qn("w:br"))
    br_elem.set(qn("w:type"), "page")
    # Añadir texto después
    run2 = para.add_run("Texto después del salto")
    return doc


def _make_doc_without_page_break() -> DocxDocument:
    """Documento normal sin saltos de página."""
    doc = DocxDocument()
    doc.add_paragraph("Párrafo uno sin salto de página.")
    doc.add_paragraph("Párrafo dos sin salto de página.")
    return doc


def _make_doc_with_table() -> DocxDocument:
    """Documento con una tabla 2x2."""
    doc = DocxDocument()
    doc.add_paragraph("Introducción.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].paragraphs[0].text = "Celda 1,1"
    table.rows[0].cells[1].paragraphs[0].text = "Celda 1,2"
    table.rows[1].cells[0].paragraphs[0].text = "Celda 2,1"
    table.rows[1].cells[1].paragraphs[0].text = "Celda 2,2"
    return doc


class TestCollectAllParagraphs:

    def test_returns_3_tuples(self):
        from app.services.correction import _collect_all_paragraphs
        doc = _make_doc_without_page_break()
        paras = _collect_all_paragraphs(doc)
        assert len(paras) >= 2
        for item in paras:
            assert len(item) == 3, "Debe retornar 3-tuple (text, location, has_page_break)"

    def test_no_page_break_detected_in_normal_doc(self):
        from app.services.correction import _collect_all_paragraphs
        doc = _make_doc_without_page_break()
        paras = _collect_all_paragraphs(doc)
        has_pb_flags = [pb for _, _, pb in paras]
        assert not any(has_pb_flags), "Documento sin saltos no debe tener flags True"

    def test_page_break_detected(self):
        from app.services.correction import _collect_all_paragraphs, _has_internal_page_break
        doc = _make_doc_with_page_break()
        paras = _collect_all_paragraphs(doc)
        has_pb_flags = [pb for _, _, pb in paras]
        assert any(has_pb_flags), "Debe detectar el salto de página interno"

    def test_table_paragraphs_no_page_break(self):
        from app.services.correction import _collect_all_paragraphs
        doc = _make_doc_with_table()
        paras = _collect_all_paragraphs(doc)
        # Los párrafos de tabla deben tener location "table:..."
        table_paras = [(t, loc, pb) for t, loc, pb in paras if loc.startswith("table:")]
        assert len(table_paras) == 4  # 2x2 = 4 celdas
        # Ninguna celda de tabla debe tener has_page_break=True
        assert not any(pb for _, _, pb in table_paras)

    def test_table_location_format(self):
        from app.services.correction import _collect_all_paragraphs
        doc = _make_doc_with_table()
        paras = _collect_all_paragraphs(doc)
        table_locations = [loc for _, loc, _ in paras if loc.startswith("table:")]
        # Formato esperado: "table:T:R:C:P"
        for loc in table_locations:
            parts = loc.split(":")
            assert len(parts) == 5, f"Location inválida: {loc}"
            assert all(p.isdigit() for p in parts[1:])

    def test_body_location_format(self):
        from app.services.correction import _collect_all_paragraphs
        doc = _make_doc_without_page_break()
        paras = _collect_all_paragraphs(doc)
        body_paras = [(t, loc, pb) for t, loc, pb in paras if loc.startswith("body:")]
        assert len(body_paras) >= 2
        for _, loc, _ in body_paras:
            parts = loc.split(":")
            assert len(parts) == 2
            assert parts[1].isdigit()


class TestHasInternalPageBreak:

    def test_detects_page_break_in_paragraph(self):
        from app.services.correction import _has_internal_page_break
        doc = _make_doc_with_page_break()
        para = doc.paragraphs[0]
        assert _has_internal_page_break(para) is True

    def test_no_false_positive_on_normal_paragraph(self):
        from app.services.correction import _has_internal_page_break
        doc = _make_doc_without_page_break()
        for para in doc.paragraphs:
            assert _has_internal_page_break(para) is False

    def test_no_false_positive_on_line_break(self):
        """w:br sin type='page' no debe detectarse como salto de página."""
        from app.services.correction import _has_internal_page_break
        doc = DocxDocument()
        para = doc.add_paragraph("Línea uno")
        run = para.runs[0]
        # Añadir w:br SIN type="page" (salto de línea)
        br_elem = etree.SubElement(run._r, qn("w:br"))
        # No añadir w:type — es salto de línea, no de página
        assert _has_internal_page_break(para) is False
